# coding: utf-8

import os
import base64
import tempfile
import sys
import re
import time as tm
from queue import Empty
from string import octdigits
from pathlib import Path
from datetime import datetime
from copy import copy

# Streamlit & path
import streamlit as st

# Excel control
import xlrd
import openpyxl
from openpyxl import Workbook, load_workbook
from openpyxl.utils.dataframe import dataframe_to_rows
from openpyxl.utils import get_column_letter
from openpyxl.styles import Font, Color, Alignment, PatternFill
from io import BytesIO

# Word file control
import docx
from docx import Document

# CSV file control
import csv 

# pdf/x contorl
from pdfminer3.pdfinterp import PDFResourceManager, PDFPageInterpreter
from pdfminer3.converter import PDFPageAggregator
from pdfminer3.pdfpage import PDFPage
from pdfminer3.layout import LAParams, LTTextContainer
import math

# Dataframe
import pandas as pd
import numpy as np


####################
from PIL import Image, ImageDraw
import easyocr
from pdf2image import convert_from_path
import pdf2image





# Initial screen for input the security code
def security_code_input():
    user_placeholder = st.empty()
    user = user_placeholder.text_input("User-name").upper()
    code_placeholder = st.empty()
    code = code_placeholder.text_input("Pass-code")

    return user, user_placeholder, code, code_placeholder


# Setting recall with mode
def mode_select(mode):
    # ---- Mashin system store style ----
    if mode == "jp_store":
        content_data = [
            "DEPT", "NOS", "", "IMPA", "ARTICLE",
            "QTY", "UNIT", "S_U_PRICE", "S_AMOUNT",
            "SUPPLIER", "*","P_U_PRICE", "P_AMOUNT",
            "REMARK","DEPT_CD", "GROUP", "SUPPLIER_CD"
        ] 
        # content identify
        order = [
            "dept_col",
            "no_col",
            "page_col",
            "impa_col",
            "article_col",
            "qty_col",
            "unit_col",
            "suprice_col",
            "samount_col",
            "empty",
            "empty",
            "puprice_col",
            "pamount_col",
            "remark_col",
            "auto_remark_col"
        ]
        # page number col specify location
        page_col = 2
        # column width identify
        column_widths = [
            12, 5, 0.38, 8.88, 39,
            8, 5, 13, 13,
            13, 3.5, 13, 13,
            36.63, 9.38, 9.38, 9.38
         ]
        font = "Arial"
        font_size = 11
        zoom = 80

    # ---- Mashin system spare style ----
    if mode == "jp_spare":
        # header formula
        content_data = [
            "DEPT", "NOS", "", "UNIT / MODEL", "ARTICLE",
            "QTY", "UNIT", "S_U_PRICE", "S_AMOUNT",
            "SUPPLIER", "*","P_U_PRICE", "P_AMOUNT",
            "REMARK","DEPT_CD", "GROUP", "SUPPLIER_CD"
        ] 
        # content identify
        order = [
            "dept_col",
            "no_col",
            "page_col",
            "impa_col",
            "article_col",
            "qty_col",
            "unit_col",
            "suprice_col",
            "samount_col",
            "empty",
            "empty",
            "puprice_col",
            "pamount_col",
            "remark_col",
            "auto_remark_col"
        ]
        # page number col specify location
        page_col = 2
        # column width identify
        column_widths = [
            12, 5, 0.38, 39, 39,
            8, 5, 13, 13,
            13, 3.5, 13, 13,
            36.63, 9.38, 9.38, 9.38
         ]
        font = "Arial"
        font_size = 11
        zoom = 80

    # ---- Other style / Netherlands sample ----
    if mode == "Netherlands":
        # header formula
        content_data = [
            "NOS", "IMPA", "ARTICLE",
            "QTY", "UNIT", "S_U_PRICE", "S_AMOUNT",
            "","REMARK","DEPT"
        ] 
        # content identify
        order = [
            "no_col",
            "impa_col",
            "article_col",
            "qty_col",
            "unit_col",
            "suprice_col",
            "samount_col",
            "page_col",
            "remark_col",
            "dept_col",
            "auto_remark_col"

        ]
        # page number col specify location
        page_col = 7
        # column width identify
        column_widths = [
            5.86, 13.30, 72.00, 6.57,
            8.57, 10.43, 11.43, 0.33,
            40.00, 8.43, 40.00
        ]
        font = "Arial"
        font_size = 12
        zoom = 80


    # ---- Other style / BUSAN sample ----
    if mode == "BUSAN":
        # header formula
        content_data = [
            "","DEPT", "NOS", "IMPA",
            "ARTICLE", "QTY", "UNIT", "S_U_PRICE",
            "S_AMOUNT","REMARK"
        ] 
        # content identify
        order = [
            "page_col",
            "dept_col",
            "no_col",
            "impa_col",
            "article_col",
            "qty_col",
            "unit_col",
            "suprice_col",
            "samount_col",
            "remark_col",
            "auto_remark_col"
        ]
        # page number col specify location
        page_col = 0
        # column width identify
        column_widths = [
            6.0, 3.63, 3.36, 5.88,
            27.75, 3.38, 3.38, 8.13,
            8.13, 10.13, 40.00
        ]
        font = "Arial Narrow"
        font_size = 9
        zoom = 100

    return content_data, order, page_col, column_widths, font, font_size, zoom


# String shape for uploaded file
def filenameshape(file):
    file_name = ""
    file_name = file.name.upper()
    file_name = file_name.replace(".XLSX", "")
    file_name = file_name.replace(".XLS", "")
    file_name = file_name.replace(".DOCX", "")
    file_name = file_name.replace(".DOC", "")###
    file_name = file_name.replace(".CSV", "")
    file_name = file_name.replace(".PDF", "")###
    file_name = file_name.strip()
    file_name = file_name.replace("　", "")
    file_name = file_name.replace(".", "")
    file_name = file_name.replace("-", "")
    file_name = file_name.replace("_", "")
    file_name = file_name.replace(",", "")
    file_name = file_name.replace("’", "")
    file_name = file_name.replace("+", "")
    file_name = file_name.replace(";", "")
    file_name = file_name.replace("；", "")
    file_name = file_name.replace("[", "")
    file_name = file_name.replace("]", "")

    return file_name


# Sort & listed by file type
@st.cache_data(show_spinner=False)
def data_sorting(uploaded_files,file, all_sht_flag):
    sheet_data = []

    # ---- Convert from xls to xlsx ----
    def xls_to_xlsx(uploaded_file, all_sht_flag):
        xls_data = uploaded_file.read()
        workbook = xlrd.open_workbook(file_contents=xls_data, on_demand=True)
        new_workbook = Workbook()
        if all_sht_flag:
            # erase sheet(0)
            default_sheet = new_workbook.active
            new_workbook.remove(default_sheet)
            # content obtain
            for sheet_name in workbook.sheet_names():
                sheet = workbook.sheet_by_name(sheet_name)
                new_sheet = new_workbook.create_sheet(title=sheet_name)
                for row in range(sheet.nrows):
                    row_data = sheet.row_values(row)
                    new_sheet.append(row_data)

        elif not all_sht_flag:
            active_sheet_name = workbook.sheet_names()[0]
            active_sheet = workbook.sheet_by_name(active_sheet_name)
            new_sheet = new_workbook.active
            new_sheet.title = active_sheet_name
            for row in range(active_sheet.nrows):
                row_data = active_sheet.row_values(row)
                new_sheet.append(row_data)

        temp_xlsx_path = "temp.xlsx"
        new_workbook.save(temp_xlsx_path)

        return temp_xlsx_path


    # ---- File read process from xlsx ----
    def read_xlsx_file(file_path, all_sht_flag):
        workbook = openpyxl.load_workbook(file_path)
        data = []
        # All sht process
        if all_sht_flag:
            for sheet_name in workbook.sheetnames:
                if sheet_name != "Statement Account": # Special case for NL
                    sheet = workbook[sheet_name]
                    sheet_data = []
            
                    # sht name & data obtian 
                    sheet_data.append([sheet_name] + [None] * (sheet.max_column - 1))
                    for row in sheet.iter_rows(values_only=True):
                        if sheet.sheet_state == 'visible':
                            sheet_data.append(row)
                    data.append(sheet_data)

        # if only target sht
        elif not all_sht_flag:
            active_sheet = workbook.active
            sheet_data = []
        
            # sht name & data obtian 
            sheet_data.append([active_sheet.title] + [None] * (active_sheet.max_column - 1))
            for row in active_sheet.iter_rows(values_only=True):
                sheet_data.append(row)
            data.append(sheet_data)

#        st.write(data)
        return data

    
    # ---- Convert from xls to xlsx ----
    if str(file.name).lower().endswith("xls"):
        if all_sht_flag:
            temp_file_path = xls_to_xlsx(file, all_sht_flag=True)
        elif not all_sht_flag:
            temp_file_path = xls_to_xlsx(file, all_sht_flag=False)
        sheet_data = read_xlsx_file(temp_file_path, all_sht_flag)
        file_type = "xls"


    # ---- Process for xlsx ----
    elif str(file.name).lower().endswith("xlsx"):
        temp_file_path = "uploaded_file.xlsx"
        with open(temp_file_path, "wb") as f:
            f.write(file.getvalue())
        sheet_data = read_xlsx_file(temp_file_path, all_sht_flag)
        file_type = "xlsx"


    # ---- Process for word docx ----
    elif str(file.name).lower().endswith("docx"):
        text_data = [p.text for p in Document(file).paragraphs]
        table = Document(file).tables[0]
        table_data = []
        for row in table.rows:
            row_data = [cell.text for cell in row.cells]
            table_data.append(row_data)

        # reflect to excel
        pd.DataFrame({'Text': text_data}).to_excel(
                                              'output.xlsx',
                                              index=False,
                                              sheet_name='TextData'
                                          )
        pd.DataFrame(table_data).to_excel(
                                     'output.xlsx',
                                     index=False,
                                     sheet_name='TableData',
                                     startrow=len(text_data)+2
                                 )
        sheet_data = read_xlsx_file("output.xlsx", all_sht_flag)
        file_type = "docx"


    # Process for word doc
    # under construnction


    # ---- Process for pdf/x ----
    # under construnction


    # ---- Process for csv ----
    elif file.type == 'text/csv':
        sheet_data = pd.read_csv(file, header=None)
        file_type = "csv"


    # ---- Exception handling / if not excel file ----
    else:
        raise ValueError("Unsupported file format: {}".format(file.name))


    # ---- Find and erase empty columns ----
    if file_type != "csv":
        empty_columns = []
        processed_data = []
        for column_index, column in enumerate(zip(*sheet_data)):
            if all(cell is None for cell in column) or all(cell == "" for cell in column):
                empty_columns.append(column_index)
        for row in sheet_data:
            processed_row = [cell for column_index, cell in enumerate(row) if column_index not in empty_columns]
            processed_data.append(processed_row)


#    st.write(sheet_data)####
    return sheet_data,file_type


# Data content identify and re-arrange
@st.cache_data(show_spinner=False)
def process_data(sheet_data, combined_data, file_type, file_name, order, page_col, vessel_name, mode):
    # ---- Load writting pattern CSV ---- 
    # A list of field names to define the structure of the CSV file
    # store pattern list load
    patternlist = [
        ['VESSELS NAME'         , 'PORT', 'DEPT'      , 'NOS'       , 'IMPA'               , 'ARTICLE'                  , 'QTY'                , 'UNIT', 'S U PRICE' , 'S_AMOUNT'    , 'REMARK'],
        ['VESSELS NAME'         , 'PORT', 'DEPT'      , 'NOS'       , 'IMPA'               , 'ARTICLE'                  , 'QTY'                , 'UNIT', 'SU_PRICE'  , 'S_AMOUNT'    , 'REMARK'],
        ['VESSEL NAME'          , ''    , 'DEPARTMENT', 'ITEMS NO'  , 'CODE'               , 'DISCRIPTION OF GOODS'     , 'QUANITY REQUIRED'   , 'UOM' , 'UNIT PRICE', 'AMOUNT'      , 'REMARKS'],
        ['SHIP S NAME'          , ''    , ''          , 'ITEM NO'   , 'IMPA  ISSA  PART NO', 'DESCRIPTION OF GOOD'      , 'REQUISITION STORES' , ''    , 'U  PRICE'  , 'TOTAL'       , 'REMERKS'],
        ['SHIPS NAME'           , ''    , ''          , 'NO'        , 'CODE PART NO'       , 'DISCRIPTION OF GOOD'      , 'REQUISITION STORE'  , ''    , 'U PRICE'   , 'TOTAL AMOUNT', 'REMERK'],
        ['SHIP NAME'            , ''    , ''          , 'RUNNING NO', 'CODEPARTNO'         , 'PART NAME AND DESCRIPTION', 'SHIP S REQUEST Q TY', ''    , 'UPRICE'    , 'T  PRICE'    , 'REASON FOR REQUEST'],
        ['M V'                  , ''    , ''          , 'S NO'      , 'OR IMPA NO'         , 'PART NAME AND DISCRIPTION', 'SHIP S REQUESTQ TY' , ''    , 'COSTS'     , 'T PRICE'     , 'ADDITIONAL INFO'],
        ['MV'                   , ''    , ''          , 'SNO'       , 'PARTS OR PLATE NO'  , 'PART S DESCRIPTION'       , 'SHIP S REQUEST QTY' , ''    , 'COST'      , 'AMT'         , 'ADDITIONAL INFORMATION'],
        ['THE MASTER & OWNER OF', ''    , ''          , 'STOCK ID'  , 'IMPA  PARTS NO'     , 'PART S DISCRIPTION'       , 'SHIP S REQUESTQTY'  , ''    , 'PRICES'    , 'S AMOUNT'    , ''],
        ['VESSEL'               , ''    , ''          , 'PARTS NO'  , 'IMPA PARTS NO'      , 'PARTS DISCRIPTION'        , 'REQUISITION'        , ''    , 'PRICE'     , ''            , ''],
        ['RE  VESSEL'           , ''    , ''          , 'ITEM'      , 'IMPA CODE'          , 'ITEM DESCRIPTION'         , 'QUANTITY'           , ''    , 'PRICE $'   , ''            , ''],
        ['RE VESSEL'            , ''    , ''          , 'PART NO'   , 'PARTS NO OR SIZE'   , 'ITEM DISCRIPTION'         , 'REQUIRED'           , ''    , 'SALE'      , ''            , ''],
        [''                     , ''    , ''          , ''          , 'SIZE  IMPA CODE'    , 'SPARE PARTS NAME'         , 'REQUEST'            , ''    , ''          , ''            , ''],
        [''                     , ''    , ''          , ''          , 'SIZE IMPA CODE'     , 'SPARE PART NAME'          , 'REQ Q TY'           , '', '', '', ''],
        [''                     , ''    , ''          , ''          , 'IMPA CODE NO'       , 'REQUISITION ITEMS'        , 'REQ QTY'            , '', '', '', ''],
        [''                     , ''    , ''          , ''          , 'SIZE   IMPA CODE'   , 'REQUISITION ITEM'         , 'REQQTY'             , '', '', '', ''],
        [''                     , ''    , ''          , ''          , 'PART NO'            , 'REQUIRED ITEMS'           , 'QTTY'               , '', '', '', ''],
        [''                     , ''    , ''          , ''          , 'PARTS NO'           , 'REQUIRED ITEM'            , 'Q TTY'              , '', '', '', ''],
        [''                     , ''    , ''          , ''          , ''                   , 'REQUEST ITEMS'            , 'Q TY'               , '', '', '', ''],
        [''                     , ''    , ''          , ''          , ''                   , 'REQUEST ITEM'             , 'QUANTITY REQUIRED'  , '', '', '', ''],
        [''                     , ''    , ''          , ''          , ''                   , 'NAME OF PARTS'            , 'QT Y'               , '', '', '', ''],
        [''                     , ''    , ''          , ''          , ''                   , 'NAME OF PART'             , 'REQ'                , '', '', '', ''],
        [''                     , ''    , ''          , ''          , ''                   , 'DESCRIPTION'              , 'RQD'                , '', '', '', ''],
        [''                     , ''    , ''          , ''          , ''                   , 'DISCRIPTION'              , 'SHIP S'             , '', '', '', ''],
        [''                     , ''    , ''          , ''          , ''                   , 'DESCRIPTIONS'             , 'QTT'                , '', '', '', ''],
        [''                     , ''    , ''          , ''          , ''                   , 'DISCRIPTIONS'             , 'REQUIRED QUANTITY'  , '', '', '', ''],
        [''                     , ''    , ''          , ''          , ''                   , 'A R T I C L E S'          , '', '', '', '', ''],
        [''                     , ''    , ''          , ''          , ''                   , 'A R T I C L E'            , '', '', '', '', ''],
        [''                     , ''    , ''          , ''          , ''                   , 'PARTS NAME'               , '', '', '', '', ''],
        [''                     , ''    , ''          , ''          , ''                   , 'PART NAME'                , '', '', '', '', ''],
        [''                     , ''    , ''          , ''          , ''                   , 'ARTICLES'                 , '', '', '', '', ''],
        [''                     , ''    , ''          , ''          , ''                   , 'DESCRIPTION OF GOODS'     , '', '', '', '', ''],
        [''                     , ''    , ''          , ''          , ''                   , 'ITEMS'                    , '', '', '', '', ''],
        [''                     , ''    , ''          , ''          , ''                   , 'ITEM'                     , '', '', '', '', ''],
        [''                     , ''    , ''          , ''          , ''                   , 'ITEM'                     , '', '', '', '', '']
    ]
    df_pattern = pd.DataFrame(patternlist)
    article_row_count = len(df_pattern.iloc[:, 5].dropna()) # pattern row count
    df_pattern_article = df_pattern.iloc[0:article_row_count, 5]

    # spare pattern list load
    spare_patternlist = [
        ["NAME OF EQUIPMENT" ,"MAKERS"                                     ,"MODEL TYPE","DRAW NO  PART NO"],
        ["NAME OF EQUIPMENT" ,"MAKERS"                                     ,"MODEL TYPE","DRAW NO  PART NO"],
        ["MACHINERYS NAME"   ,"MAKER"                                      ,"MODELTYPE" ,"DRAW NO PART NO"],
        ["NAME OF MACHINERYS","M A K E R S"                                ,"MODELS"    ,"DWG NO PART NO"],
        ["NAME OF MACHINERY" ,"M A K E R"                                  ,"MODEL"     ,"SIZE MAKER P NO"],
        ["MACHINERY NAME"    ,"COMMENTS"                                   ,"M O D E L" ,"PART  CODE"],
        ["MACHINES NAME"     ,"COMMENT"                                    ,"T Y P E"   ,"PART CODE"],
        ["MACHINE NAME"      ,"SIZE  STANDARD  MATERIAL"                   ,"TYPES"     ,"MATERIALS"],
        ["ENGINES TYPE"      ,"SIZE  STANDARD  MATERIAL OTHER INFORMATION" ,"TYPE"      ,"MATERIAL"],
        ["ENGINE TYPE"       ,"SIZE  STANDARD  MATERIAL  OTHER INFORMATION",""          ,"PARTS NO"],
        ["EQUIPMENTS"        ,""                                           ,""          ,"PART NO"],
        ["EQUIPMENT"         ,""                                           ,""          ,"CODE NO"],
        ["MACHINERYS"        ,""                                           ,""          ,"CORD NO"],
        ["MACHINERY"         ,""                                           ,""          ,"EDITION"],
        [""                  ,""                                           ,""          ,"SPECS"],
        [""                  ,""                                           ,""          ,"SPEC"],
        [""                  ,""                                           ,""          ,"ID NO"],
        [""                  ,""                                           ,""          ,"SIZE"],
        [""                  ,""                                           ,""          ,"TYPE"],
        [""                  ,""                                           ,""          ,"TYPE"]
    ]
    df_spare_pattern = pd.DataFrame(spare_patternlist)

    # dept pattern list load
    deptlist = [
        ['DECK'       , 'ENGINE'    , 'STATIONERY', 'STEWARD'           , 'MEDICAL' , 'CHARGE'],
        ['CHART'      , 'PIPE'      , 'STATIONARY', 'SETWARD'           , 'MED-ADD' , 'EXPENSE'],
        ['PUBLICATION', 'ELECTRICAL', 'STA-ADD'   , 'STAWARD'           , 'MEDICINE', 'EXPENSIVE'],
        ['FLAG'       , 'ELECTRIC'  , ''          , 'STW-ADD'           , 'SURGICAL', 'SEAFREIGHT'],
        ['RADIO'      , 'CHEMICAL'  , ''          , 'GALLEY'            , 'HOSPITAL', 'AIRCARGO'],
        ['SAFETY'     , 'ENG ADD'   , ''          , 'PURSER'            ,''         , 'TRANSPORT'],
        ['BRIDGE'     , 'BEARING'   , ''          , 'CATERING'          , ''        , 'TRANSPORTATION'],
        ['LSA'        , 'ANODE'     , ''          , 'CABIN'             , ''        , ''],
        ['FFE'        , 'ELE'       , ''          , 'GALLEY AND STEWARD', ''        , ''],
        ['LSAFFE'     , 'ENG'       , ''          , ''                  , ''        , ''],
        ['SYMBOL'     , '', '', '', '', ''],
        ['PAINT'      , '', '', '', '', ''],
        ['SOPEP'      , '', '', '', '', ''],
        ['NAVIGATION' , '', '', '', '', ''],
        ['EARLY'      , '', '', '', '', ''],
        ['TRANSCEIVER', '', '', '', '', ''],
        ['OTHER'      , '', '', '', '', ''],
        ['OTHERS'     , '', '', '', '', ''],
        ['LOGBOOK'    , '', '', '', '', ''],
        ['LOGBOOKS'   , '', '', '', '', ''],
        ['LOG BOOK'   , '', '', '', '', ''],
        ['LOG BOOKS'  , '', '', '', '', ''],
        ['COVID'      , '', '', '', '', ''],
        ['CARGO HOLD' , '', '', '', '', ''],
        ['IMO SYMBOL' , '', '', '', '', ''],
        ['SYMBOL'     , '', '', '', '', ''],
        ['DIGITAL'    , '', '', '', '', ''],
        ['SANITARY'   , '', '', '', '', ''],
        ['GYM'        , '', '', '', '', ''],
        ['LIFE SAVING APPLIANCES', '', '', '', '', ''],
        ['DECKSAFETY' , '', '', '', '', ''],
        ['DECK SAFETY', '', '', '', '', ''],
        ['CALIBRATION', '', '', '', '', ''],
        ['CALIBRATION', '', '', '', '', '']
    ]

    # Advance setting
    data_index = sheet_key = page_col_value = 0 # for can not obtain actual col
    vessel_flag = False
    port_name = ""


    # ---- Data obtain process ----
    for data in iter(sheet_data):
        none_sht_flag = False
        # set df
        if file_type == 'csv':
            df = sheet_data
        else:
            processed_data = list(data)
            df = pd.DataFrame(processed_data)
#        st.write(df)####

        # advance content control
        df_content = df.applymap(lambda x: x.upper() if isinstance(x, str) else x) # all strings uppercase
        df_content.iloc[0:2] = df_content.iloc[0:2].applymap(lambda x: x.replace("PRICE_ERROR", "") if isinstance(x, str) else x)
        df_content.iloc[0:2] = df_content.iloc[0:2].applymap(lambda x: x.replace("SUPPLIER_CD2", "") if isinstance(x, str) else x)
        df_content.iloc[0:2] = df_content.iloc[0:2].applymap(lambda x: x.replace("SUPPLIER2", "") if isinstance(x, str) else x)
        df_content.iloc[0:2] = df_content.iloc[0:2].applymap(lambda x: x.replace("P_U_PRICE2", "") if isinstance(x, str) else x)
        df_content.iloc[0:2] = df_content.iloc[0:2].applymap(lambda x: x.replace("SUPPLIER_CD3", "") if isinstance(x, str) else x)
        df_content.iloc[0:2] = df_content.iloc[0:2].applymap(lambda x: x.replace("SUPPLIER3", "") if isinstance(x, str) else x)
        df_content.iloc[0:2] = df_content.iloc[0:2].applymap(lambda x: x.replace("P_U_PRICE3", "") if isinstance(x, str) else x)
        df_content.replace("　", " ",regex=True, inplace=True)
        df_content.replace("_", " ",regex=True, inplace=True)
        df_content.replace("\(", " ",regex=True, inplace=True)
        df_content.replace("\)", " ",regex=True, inplace=True)
        df_content.replace(":", " ",regex=True, inplace=True)
        df_content.replace("：", " ",regex=True, inplace=True)
        df_content.replace(";", " ",regex=True, inplace=True)
        df_content.replace("；", " ",regex=True, inplace=True)
        df_content.replace("'", "",regex=True, inplace=True)
        df_content.replace("’", "",regex=True, inplace=True)
        df_content.replace("/", " ",regex=True, inplace=True)
        df_content.replace("\"", " ",regex=True, inplace=True)
        df_content.replace("\n", " ",regex=True, inplace=True)
        df_content = df_content.applymap(lambda x: x.strip() if isinstance(x, str) else x) # trim
        df_content = df_content.applymap(lambda x: int(x) if isinstance(x, str) and x.isdigit() else x) # str -> int
        df = df_content

        # col sontrol
        df = df.dropna(axis=1, how='all') # delete None col
        df = df.iloc[:, :80] # delete column over 80
        df.columns = range(len(df.columns)) # re-set col index
        df_col_count = len(df.columns)

        # row control
        # numeric 1 str to be None
        for i, row in df.iterrows():
            if row.count() == 1:
                if pd.to_numeric(row, errors='coerce').count() == 1:
                    df.iloc[i] = [None] * len(df.columns)

        # delete empty row with 5rows
        empty_row_count = 0
        rows_to_delete = []
        for i, row in df.iterrows():
            if row.isnull().all() or row.astype(str).str.strip().eq('').all():
                empty_row_count += 1
                if empty_row_count >= 5:
                    rows_to_delete.extend(range(i - empty_row_count + 1, i + 1))
                    # Replace the last empty row with another empty row
                    if i - empty_row_count >= 0:
                        rows_to_delete.extend(range(i - empty_row_count + 1, i))
                        empty_row_count = 1
            else:
                empty_row_count = 0
        df = df.drop(rows_to_delete)

        # delete empty row with 3rows
        none_row_count = 0
        rows_to_drop = []
        for i, row in df.iterrows():
            if all(value is None for value in row):
                none_row_count += 1
            else:
                none_row_count = 0       

            if none_row_count == 3:
                rows_to_drop.append(i-2)
        df = df.drop(rows_to_drop).reset_index(drop=True)
       
        # row re-identify
        df.index = range(len(df.index)) # re-set row index
        df_row_count = len(df)

        # data available fact check
        none_sht_flag = False
        if df_row_count < 2 and df_col_count < 2: # skip to next dats if data is empty
            none_sht_flag = True

#        st.write(str(df_col_count))####
#        st.write(str(df_row_count))####
#        st.write(df)####


        # ---- Article row & col idenfity ----
        article_row = article_col = ""
        fact_flag = False
        for str_check in df_pattern_article:
            if not fact_flag:
                indices = df.astype(str).apply(lambda x: x.str.contains(str_check, case=False))
                article_row, article_col = indices.values.nonzero() # article_row,col is "numpy.ndarray"
#                st.write(str_check)
#                st.write(article_row)####
#                st.write(article_col)####

                # data available fact check
                if article_row.size > 0:
                    none_count = 0
                    for row in range(article_row[0],df_row_count):
                        for col in range(df_col_count):
                            cell_value = df.iloc[row, col]
                            if cell_value is None:
                                none_count += 1
                    # 1st faxt check by None
                    if none_count / ((df_row_count - article_row[0]) * df_col_count) > 0.95:
                        none_sht_flag = True
                        break

                # Fact check process after fixed the article row & col
                fact_check_article_col = article_col
                if article_row.size > 0:
                    for fact_cnt in range(3): # Check up to 2 cols right
                        fact_flag = False
                        num_cnt = empty_cnt = 0
                        for i in range(article_row[0] + 1, df.shape[0]): # df.shape[0] = article_row.rowcount
                            cell_value = df.iloc[i, article_col + fact_cnt]
                            if pd.to_numeric(cell_value, errors='coerce').notnull().all(): # if it can be interpreted as a number
                                num_cnt += 1
                            elif pd.isnull(cell_value).all(): # Empty number count / True = Empty
                                empty_cnt += 1
                        
                        # Fact check result confirmation / clear condition = Empty under 90%  or  Number under 30%
                        if (empty_cnt / (df.shape[0] - (article_row[0] + 1))) < 0.9 and (num_cnt / (df.shape[0] - (article_row[0] + 1))) < 0.3:
                            fact_flag = True
                            article_col = article_col[0] + fact_cnt
                            article_row = article_row[0]
                            break

                    # backward fact check
                    if not fact_flag:
                        fact_cnt = -1 # backward place set
                        num_cnt = empty_cnt = 0
                        for i in range(article_row[0] + 1, df.shape[0]): # df.shape[0] = article_row.rowcount
                            cell_value = df.iloc[i, fact_check_article_col + fact_cnt]
                            if pd.to_numeric(cell_value, errors='coerce').notnull().all(): # if it can be interpreted as a number
                                num_cnt += 1
                            elif pd.isnull(cell_value).all(): # Empty number count / True = Empty
                                empty_cnt += 1
                        
                        # Fact check result confirmation / clear condition = Empty under 90%  or  Number under 30%
                        if (empty_cnt / (df.shape[0] - (article_row + 1))) < 0.9 and (num_cnt / (df.shape[0] - (article_row + 1))) < 0.3:
                            fact_flag = True
                            article_col = fact_check_article_col[0] + fact_cnt
                            article_row = article_row[0]
                            break


                        # ---- if can not obtain the actual col ----
                        if not fact_flag:
                            num_input = st.empty()
                            input_list_disp = st.empty()
                            num_input_value = ""
                            try:
                                data_index += 1
                                sht_key = f"sht_key_{data_index}"
                                num_input_value = num_input.text_input(
                                    "Please input the actual column of Description",
                                    key=sht_key,
                                    value=num_input_value,
                                )
                                input_list = df.iloc[article_row[0] - 1:article_row[-1] + 4,:]
                                input_list_disp.write(input_list)
                            except Exception as e:
                                st.error(" ")
                                sys.exit(1)

                            # after actual col input
                            if num_input_value != '':
                                fact_flag = True
                                article_col = (article_col[0], int(num_input_value))
                                article_row = article_row[0]
                                article_col = article_col[1]
                                num_input.empty()
                                input_list_disp.empty()

                elif fact_flag:
                    break

        # None sheet counterplan - skip to next data if no content
        if none_sht_flag:
            continue


        # ---- if article N/A ----
        if not article_row.size:
            input_row = st.empty()
            input_col = st.empty()
            input_list_disp_add = st.empty()
            num_input_row = ""
            num_input_col = ""
            data_index += 1
            input_list_disp_add.write(df.iloc[:20])
            num_input_row = input_row.text_input(
                "Please input the actual row of Description",
                key= f"row_key_{data_index}",
                value=num_input_row,
            )
            num_input_col = input_col.text_input(
                "Please input the actual column of Description",
                key=f"col_key_{data_index}",
                value=num_input_col,
            )

            # after article_row & col input                
            article_row = int(num_input_row) if num_input_row is not None and num_input_row != "" else None
            article_col = int(num_input_col) if num_input_col is not None and num_input_col != "" else None
            if article_row != None and article_row != "":
                input_row.empty()
            if article_col != None and article_col != "":
                input_col.empty()
                input_list_disp_add.empty()
            fact_flag = True

        # Article row & col fixed
#        st.write(article_row)####
#        st.write(article_col)####


        # ---- shape after fixed the article col - target is only article_row ----
        for col in range(df_col_count):
            string = df.iloc[article_row, col]
            if string is not None:
                string = str(string)
                replaced_string = re.sub(r'\d+', '', string)
                replaced_string2 = replaced_string.replace(".", "")
                replaced_string3 = replaced_string2.strip()
                df.iloc[article_row, col] = replaced_string3

        # add row which only 1 item in col
        target_rows = []
        new_row = [None] * df.shape[1]

        # fixed the target row
        for index in range(article_row, df.shape[0]):
            row = df.iloc[index]
            if row.count() == 1 and df.iloc[index, article_col] is not None:
                target_rows.append(index)
        if target_rows: # erase the first row
            target_rows.pop(0)

        # add empty row at target row
        for target_row in target_rows:
            target_row += 1 
            df = pd.concat([df.iloc[:target_row], pd.DataFrame([new_row], columns=df.columns), df.iloc[target_row:]], ignore_index=True)
        df = df.reset_index(drop=True)
#        st.write(df)####

 
        # ---- renew roucount ----
        df_row_count = len(df)
#        st.write(df_row_count)####


        # ---- vessel name obtain ----
        def name_shape(target_name):
            target_name = target_name.replace(".", "")
            target_name = target_name.replace("-", "")
            target_name = target_name.replace("/", "")
            target_name = target_name.replace("(", "")
            target_name = target_name.replace(")", "")
            target_name = re.sub(r'\d+', '', target_name)
            target_name = target_name.strip()
            return target_name
        
        # obtain the vessel name only one time
        if vessel_name == "":
            vessel_row = []
            vessel_col = []
            vessel_flag = False
            df_pattern_vessel = df_pattern.iloc[0:article_row_count, 0]

            # target place identify
            for str_check in df_pattern_vessel:
                if not vessel_flag and str_check != "":
                    indices = df.astype(str).apply(lambda x: x.str.contains(str_check, case=False))
                    vessel_row, vessel_col = indices.values.nonzero() # vessel_row,col is "numpy.ndarray"
                    if vessel_row.size > 0:
                        # target str place in df
                        vessel_row = vessel_row[0]
                        vessel_col = vessel_col[0]
                        
                        # str replace
                        vessel_name = df.iloc[vessel_row, vessel_col].upper()
                        target_name = vessel_name.replace(str(str_check),"")
                        vessel_name = name_shape(target_name)
                        if vessel_name == "":
                            vessel_flag = True
                        break

            # fact check for target str
            if vessel_flag:
                for i in range(4):
                    if df.iloc[vessel_row, vessel_col + i] is not None:
                        vessel_name = df.iloc[vessel_row, vessel_col + i].upper()
                        if vessel_name:
                            target_name = vessel_name.replace(str(str_check),"")
                            vessel_name = name_shape(target_name)
                        if vessel_name != "":
                            break

             # vessel name support
            if vessel_name:
                vessel_name = vessel_name.replace("MV","")
                vessel_name = vessel_name.replace("M V","")
                vessel_name = vessel_name.strip()
                vessel_name = vessel_name + " "

#        st.write(vessel_name)####


        # ---- port name obtain ----
        if port_name == "":
            port_row = []
            port_col = []
            port_flag = False
            df_pattern_port = df_pattern.iloc[0:article_row_count, 1]

            # target place identify
            for str_check in df_pattern_port:
                if not port_flag and str_check != "":
                    indices = df.astype(str).apply(lambda x: x.str.contains(str_check, case=False))
                    port_row, port_col = indices.values.nonzero() # port_row,col is "numpy.ndarray"
                    if port_row.size > 0:
                        # target str place in df
                        port_row = port_row[0]
                        port_col = port_col[0]
                        
                        # str replace
                        port_name = df.iloc[port_row, port_col].upper()
                        target_name = port_name.replace(str(str_check),"")
                        port_name = name_shape(target_name)
                        if port_name == "":
                            port_flag = True
                        break

            # fact check for target str
            if port_flag:
                for i in range(4):
                    if df.iloc[port_row, port_col + i] is not None:
                        port_name = df.iloc[port_row, port_col + i].upper()
                        if port_name:
                            target_name = port_name.replace(str(str_check),"")
                            port_name = name_shape(target_name)
                        if port_name != "":
                            break

             # vessel name support
            if port_name:
                port_name = port_name.replace("PORT NAME","")
                port_name = port_name.replace("PORT","")
                port_name = port_name.strip()

#        st.write(port_name)####





        # ---- spare part name obtain ----
        # shape method
        def spare_shape(target_name):
            target_name = target_name.replace("/", " ")
            target_name = target_name.replace("-", " ")
            target_name = target_name.strip()
            return target_name

        # obtain the spare part name only one time
        def spare_content(position):
            spare_name = ""
            spare_row = []
            spare_col = []
            spare_flag = False
            spare_df = df.iloc[2:article_row + 4, :]
            df_pattern_spare = df_spare_pattern.iloc[0:len(spare_df), position]
            
            # target place identify
            for str_check in df_pattern_spare:
                if not spare_flag and str_check != "":
                    indices = spare_df.astype(str).apply(lambda x: x.str.contains(str_check, case=False))
                    spare_row, spare_col = indices.values.nonzero() # spare_row,col is "numpy.ndarray"
                    if spare_row.size > 0:
                        # target str place in df
                        spare_row = spare_row[0]
                        spare_col = spare_col[0]
                        
                        # str replace
                        spare_name = spare_df.iloc[spare_row, spare_col].upper()
                        target_name = spare_name.replace(str(str_check),"")
                        spare_name = spare_shape(target_name)
                        if spare_name == "":
                            spare_flag = True
                            break

            # fact check for target str
            if spare_flag:
                for i in range(4):
                    if spare_df.iloc[spare_row, spare_col + i] is not None:
                        spare_name = spare_df.iloc[spare_row, spare_col + i].upper()
                        if spare_name:
                            target_name = spare_name.replace(str(str_check),"")
                            spare_name = spare_shape(target_name)
                        if spare_name != "":
                            break               
            return spare_name

        # spare content obtain call def
        if mode == "jp_spare":
            spare_identify = ""

            # MACHINE
            position = 0
            machine_name = spare_content(position)
            if machine_name != "":
                machine_name = machine_name + ", "
#            st.write("machine " + machine_name)####

            # maker
            position = 1
            spare_maker = spare_content(position)
            if spare_maker != "":
                spare_maker = spare_maker + ", "
#            st.write("maker " + spare_maker)####

            # type
            position = 2
            spare_type = spare_content(position)
            if spare_type != "":
                spare_type = spare_type + " "
#            st.write("type " + spare_type)####

            # drawing
            position = 3
            spare_drw = spare_content(position)
            if spare_drw.startswith(".") and spare_drw != "":
                spare_drw = spare_drw.lstrip(".")
#            st.write("drw" + spare_drw)####

            # info combine
            spare_identify = machine_name + spare_maker + spare_type + spare_drw
            if spare_identify.startswith(".") and spare_identify != "":
                spare_identify = spare_identify.lstrip(",")
#            st.write(spare_identify)####


        # ---- prepare for Other content row & col obtain ----
        # target str adjust
        if fact_flag:
            for col in df.columns:
                df.loc[article_row, col] = (
                    str(df.loc[article_row, col])
                    .replace(".", " ")
                    .replace("/", " ")
                    .replace("-", " ")
                    .upper() 
                    .strip("\n").strip() # remove leading newline
                    .strip()
                )
        
        # Set the reader content
        df_pattern_cont = [
            ["page_col","1"],
            ["dept_col","2"],
            ["no_col","3"],
            ["impa_col","4"],
            ["article_col","5"],
            ["qty_col","6"],
            ["unit_col","7"],
            ["suprice_col","8"],
            ["samount_col","9"],
            ["remark_col","10"],
            ["empty", "-1"]  # empty for blank column
        ]

        # Other content row & col obtain
        def content_position(col_number,pattern_flag):
            # df pattern set
            cont_pattern = df_pattern.iloc[:, int(col_number)].dropna()
            if int(col_number) >= df_pattern.shape[1]:
                return 99
            cont_row_count = len(cont_pattern[cont_pattern != ""]) # rowcount without ""
            cont_pattern = df_pattern.iloc[0:cont_row_count, int(col_number)]
            cont_col = 99 # initial value - Avoid duplication with list 0
            found_col = False

            # content df set
            content_df = []
            content_df = df.iloc[article_row:df.shape[0], 0:df.shape[1]] # dataframe fm article_row to row count

            # Content col identify
            for str_check in cont_pattern:
                if cont_col == 0:  # Skip loop if cont_col is already assigned
                    break

                for col in content_df.columns:
                    if str_check.strip() != "":
                        # Exact string match  but.uppercase or lowercase doesn't matter
                        indices = content_df[col].astype(str).str.contains(str_check, case=False, regex=False) & (content_df[col].astype(str) == str_check)
                        if indices.any():
                            if not found_col:
                                cont_col = col
                            # Skip the remaining rows for this column
                            elif col != cont_col:
                                break
            
            # If cont_col is still 99(=None), search in the next row
            if cont_col == 99:
                # set the roop count
                roop_cnt = 0
                if (df_row_count - article_row) < 4:
                    roop_cnt = df_row_count - article_row
                else:
                    roop_cnt = 4

                # start to obtain 
                for i in range(1,roop_cnt):
                    if i < content_df.shape[0]: 
                        for col in content_df.columns:
                            if str_check.strip() != "":
                                value = content_df.loc[article_row + i, col]
                                if value is not None:
                                    indices = content_df[col].astype(str).str.contains(str_check, case=False, regex=False) & (content_df[col].astype(str) == str_check)
                                    if indices.any():
                                        found_col = True
                                        cont_col = col
                                        break


            # ---- special case for QTY & Unit separate ----
            if col_number == "6":
                qty_num = 0
                qty_mix = 0
                qty_str = 0
                column_name = 0
                # check qty content - Numbers Only,Numbers & Letters or Letters Only
                for i in range(article_row+1, df_row_count):
                    row = df.iloc[i]
                    qty_content = row[cont_col]
                    if qty_content is not None:
                        if isinstance(qty_content, int) or isinstance(qty_content, float): # natural num
                            qty_num += 1
                        elif isinstance(qty_content, str) and qty_content.isdigit(): # possible to convert from str to int
                            qty_num += 1
                        elif isinstance(qty_content, str) and any(char.isdigit() for char in qty_content): # str & num mixed
                            qty_mix += 1
                        elif isinstance(qty_content, str): # natural str
                            qty_str += 1

                # identify the main qty content
                max_qty = max(qty_num, qty_mix, qty_str)
#                st.write(max_qty)####

                # correspondence with qty status
                # Most popular item is int / qty_col = cont_col
                if max_qty == qty_num:
                    return cont_col,pattern_flag

                # Most popular item is int & str mixed - separate qty & make new col for unit
                elif max_qty == qty_mix:
                    # New column added at df.col.count
                    new_column_values = [''] * df_row_count 
                    column_name = len(df.columns)
                    df[column_name] = new_column_values

                    # sparate qty & unit
                    for i in range(article_row + 1, df_row_count):
                        row = df.iloc[i]
                        unit_cont = ''.join([char for char in str(row[cont_col]) if char.isalpha()])
                        if unit_cont != "" and df.at[i,cont_col] is not None:
                            # unit content adjust
                            unit_cont = str(unit_cont).upper()
                            unit_cont = str(unit_cont).replace(" ", "")
                            unit_cont = str(unit_cont).strip()
                            if unit_cont[-1] == "S":
                                unit_cont = unit_cont[:-1]
                            if unit_cont == "" or unit_cont == "Null":
                                unit_cont = None
                            df.at[i,column_name] = unit_cont

                            # qty content adjust
                            qty_cont = str(row[cont_col]).upper()
                            qty_cont = qty_cont.replace(unit_cont,"")
                            qty_cont = qty_cont.replace(" ", "")
                            qty_cont = qty_cont.replace("S", "")
                            qty_cont = qty_cont.strip()

                            # qty content set
                            df.at[i,cont_col] = qty_cont
                    
                    pattern_flag = column_name
                    return cont_col,pattern_flag

                # if Most popular item is str
                elif max_qty == qty_str:
                    pass

            return cont_col,pattern_flag

        # Content col fixed method
        fixed_col = []
        pattern_flag = 99
        for pattern in df_pattern_cont:
            col_name = pattern[0]
            col_number = pattern[1]
            result, pattern_flag = content_position(col_number,pattern_flag)
            if pattern[0] == "empty" or result == 99:
                result = -1

            # append to fixed col
            if col_name == "article_col":
                fixed_col.append([col_name, article_col])
            elif col_name == "unit_col" and pattern_flag != 99:
                fixed_col.append([col_name, pattern_flag])
            else:
                fixed_col.append([col_name, result])

            # dept support
            if col_name == "dept_col":
                dept_col_value = result
#        st.write(fixed_col)####


        # ---- sort & set by mashin system format ----
        if fact_flag:
            article_row = article_row + 1

        # set col-1 If there is no corresponding item
        column_indices = []
        for col_name in order:
            col_index = -1
            for i, fixed in enumerate(fixed_col):
                if fixed[0] == col_name:
                    col_index = fixed[1]
                    break
            column_indices.append(col_index)  
#        st.write(column_indices)####

        # file name col append & info 
        sorted_data = []
        page_col_value += 1
        if page_col_value == 1:
            new_row = [None] * len(order)
            new_row[order.index("article_col")] = file_name
            sorted_data.append(new_row)

        # value None check / "" -> None
        for row in range(df_row_count):
            for col in range(len(df.columns)):
                value = df.iloc[row, col]
                if value == "" or value == "Null":
                    df.at[row, col] = None

        # sort by mashin system format
        if fact_flag:
            for row in range(article_row,df_row_count):
                row_data = []
                for col_index in column_indices:
                    if col_index != -1:
                        # value None check / "" -> None
                        value = df.iloc[row, col_index]
                        if value == "":
                            value = None
                        row_data.append(df.iloc[row, col_index])
                    else:
                        row_data.append("")

                sorted_data.append(row_data)
#        st.write(sorted_data)####


        # ---- adjust after sort by format ----
        # page num value set & no_col adjust
        for row_data in sorted_data:
            if len(row_data) > page_col:
                row_data[page_col] = page_col_value

        # sheet name info
        sht_name_value = df.iloc[0, 0]
        if isinstance(sht_name_value, int) or isinstance(sht_name_value, float):
            sht_name = "sheet1"
        else:
            sht_name = str(df.iloc[0, 0]).upper()
            sht_name = name_shape(sht_name)
            sht_name = sht_name.replace("_", "")
            sht_name = sht_name.replace("'", "")
            sht_name = sht_name.replace("’", "")
            sht_name = sht_name.replace(",", "")
            sht_name = sht_name.replace("、", "")

        if not "SHEET" in sht_name:
            if sorted_data and page_col < len(sorted_data[0]):
                if page_col_value == 1:
                    sorted_data[1][page_col] = sht_name
                elif page_col_value > 1:
                    sorted_data[0][page_col] = sht_name

        # sheet name  spare mode special case
        if mode == "jp_spare":
            for row in range(len(sorted_data)):
                # dept input support
                sorted_data[row][0] = "ENGINE"

                # spare name,type info
                if sorted_data[row][4] is not None:
                    sorted_data[row][3] = spare_identify

        # no_col adjust
        for row in range(len(sorted_data)):
            if sorted_data[row][1] is not None:
                if isinstance(sorted_data[row][1], str) and "." in sorted_data[row][1]:
                    sorted_data[row][1] = sorted_data[row][1].replace(".", "")

#        st.write(sorted_data)####


        # ---- dept formula support ----
        # dept available check
        if (dept_col_value == -1 or dept_col_value == 99) and mode != "jp_spare":
            # New dept col make
            for row in sorted_data:
                row.append(None)
            dept_col_index = len(sorted_data[0]) - 1 # for dept change flag
#            st.write(dept_col_index)####

            # dept change flag
            deptflag = True
            sorted_data[0][dept_col_index] = "na_st"
            empty_count = 0
            for row in range(len(sorted_data)-1):
                if sorted_data[row][order.index("article_col")] is None:
                    empty_count += 1
                    if empty_count > 1 and sorted_data[row+1][order.index("article_col")] is not None:
                        sorted_data[row+1][dept_col_index] = "na_st"
                        empty_count = 0
                elif sorted_data[row][order.index("article_col")] is not None:
                    empty_count = 0

            # identify the dept name at NA_ST
            for row in range(len(sorted_data)):
                sort_wor_count = len(sorted_data)
                if sorted_data[row][dept_col_index] == "na_st":
                    article_check_pluszero = ""
                    article_check_plusone = ""
                    article_check_plustwo = ""
                    hit_row_value = ""

                    # Obtain the content
                    def get_article_check(sorted_data, order, index, row, sort_wor_count):
                        if sort_wor_count > row+index:
                            article_check = sorted_data[row+index][order.index("article_col")]
                            if isinstance(article_check, str) and article_check is not None and article_check[-1] == "S":
                                article_check = article_check.replace("S", "")
                            return article_check

                    article_check_pluszero = get_article_check(sorted_data, order, 0, row, sort_wor_count)
                    article_check_plusone = get_article_check(sorted_data, order, 1, row, sort_wor_count)
                    article_check_plustwo = get_article_check(sorted_data, order, 2, row, sort_wor_count)

                    # dept name included check
                    def find_hit_row_col(deptlist, article_check):
                        hit_row = None
                        hit_col = None
                        for i, row in enumerate(deptlist):
                            for j, col in enumerate(row):
                                if isinstance(article_check, str) and col != "" and col in article_check:
                                    hit_row = i
                                    hit_col = j
                                    break
                            if hit_row is not None:
                                break
                        return hit_row, hit_col

                    # dept name included check for pluszero
                    if article_check_pluszero is not None:
                        hit_row, hit_col = find_hit_row_col(deptlist, article_check_pluszero)
                        if hit_row is not None and hit_col is not None:
                            hit_row_value = deptlist[0][hit_col]

                    # dept name included check for plusone
                    if not hit_row_value and article_check_plusone is not None:
                        hit_row, hit_col = find_hit_row_col(deptlist, article_check_plusone)
                        if hit_row is not None and hit_col is not None:
                            hit_row_value = deptlist[0][hit_col]

                    # dept name included check for plustwo
                    if not hit_row_value and article_check_plustwo is not None:
                        hit_row, hit_col = find_hit_row_col(deptlist, article_check_plustwo)
                        if hit_row is not None and hit_col is not None:
                            hit_row_value = deptlist[hit_row][hit_col]

#                    st.write(hit_row_value)####

                    # dept name set
                    sorted_data[row][order.index("dept_col")] = hit_row_value

            # dept name end-point set
            na_flag= False
            for row in range(len(sorted_data)-1, -1, -1):
                if sorted_data[row][dept_col_index] == "na_st":
                    na_flag = True
                elif na_flag and sorted_data[row][order.index("article_col")] is not None:
                    sorted_data[row][order.index("dept_col")] = "end_point"
                    na_flag = False

            # if dept is not available / try to set fm sheet name
            dept_fm_shtname = ""
            hit_row, hit_col = find_hit_row_col(deptlist, sht_name)
            if hit_row is not None and hit_col is not None:
                dept_fm_shtname = deptlist[0][hit_col]
#            st.write(dept_fm_shtname)####

            # if dept is not available / try to set fm file name
            if dept_fm_shtname == "":
                hit_row, hit_col = find_hit_row_col(deptlist, file_name)
                if hit_row is not None and hit_col is not None:
                    dept_fm_shtname = deptlist[0][hit_col]
#                st.write(dept_fm_shtname)####


            # input the dept name with actual_dept & sht_name
            actual_dept = ""
            for row in range(len(sorted_data)):
                # advance set for if
                dept_col =sorted_data[row][order.index("dept_col")]
                dept_index = sorted_data[row][dept_col_index]

                # available dept name start-point
                if dept_index == "na_st" and dept_col != "" and actual_dept == "":
                    actual_dept = sorted_data[row][order.index("dept_col")]
                    sorted_data[row][order.index("dept_col")] = actual_dept
                    sorted_data[row][dept_col_index] = ""

                # not available dept name start-point / via sht_name
                if dept_index == "na_st" and dept_col == "" and actual_dept == "":
                    sorted_data[row][order.index("dept_col")] = dept_fm_shtname 
                    actual_dept = sorted_data[row][order.index("dept_col")]
                    sorted_data[row][order.index("dept_col")] = actual_dept
                    sorted_data[row][dept_col_index] = ""

                # dept input continue
                elif dept_col == "" and actual_dept != "":
                   sorted_data[row][order.index("dept_col")] = actual_dept

                # dept input end-point 
                elif dept_col == "end_point":
                   sorted_data[row][order.index("dept_col")] = actual_dept
                   sorted_data[row+1][order.index("dept_col")] = actual_dept
                   actual_dept = ""

                # dept identify rest na_st delete
                sorted_data[row][dept_col_index] = ""
#                st.write(actual_dept)####



        # ---- other formula shape ----
        for row in range(len(sorted_data)):
            # IMPA col formula
            if sorted_data[row][order.index("impa_col")] is not None:
                sorted_data[row][order.index("impa_col")] = spare_shape(str(sorted_data[row][order.index("impa_col")]))
                sorted_data[row][order.index("impa_col")] = str(sorted_data[row][order.index("impa_col")]).replace(" ", "")

            if sorted_data[row][order.index("unit_col")] is not None:
                str_check = sorted_data[row][order.index("unit_col")]
                if str_check and len(str_check) > 0:
                    str_check = str_check.replace(".", "")
                    str_check = str_check.strip()
                    if str_check and len(str_check) > 0:
                        if str_check[-1] == "S" and str_check != "C/S":
                            sorted_data[row][order.index("unit_col")] = str_check[:-1]

            # article content shape / If there are two space to be one space
            if sorted_data[row][order.index("article_col")] is not None:
                for col in range(len(sorted_data[row])):
                    if isinstance(sorted_data[row][col], str):
                        sorted_data[row][col] = " ".join(sorted_data[row][col].split())

#        st.write(sorted_data)####


        # ---- Data extend per sheet ----
        combined_data.extend(sorted_data)
        combined_data[0][0]= "" # erase for dept-col of shee name
        max_column = max(len(row) for row in sorted_data)
        empty_row = [None] * max_column #len(sorted_data) # empty row
        for _ in range(3):  # append empty row 3 times
            combined_data.append(empty_row)

#        st.write(combined_data)####


        # ---- special case for csv ----
        if file_type == 'csv':
            break
    
    return combined_data, vessel_name, port_name


# Autoremark process
def autoremark(combined_data, df_remarkbase_1, df_remarkbase_2, order):
    # ---- advance setting ----
    row_count = len(combined_data)
    remark_content = [[] for _ in range(row_count)]

    # ---- QTY base data set ----
    qtydup_table_deck = [
                            ['dept', 'impa', 'qty', 'unit'],
                            ['DECK', '174058', '20', 'PC'],
                            ['DECK', '174276', '24', 'PC'],
                            ['DECK', '174277', '50', 'PC'],
                            ['DECK', '174291', '4', 'PC'],
                            ['DECK', '174292', '4', 'PC'],
                            ['DECK', '174295', '5', 'PC'],
                            ['DECK', '190101', '50', 'DOZ'],
                            ['DECK', '190109', '2', 'DOZ'],
                            ['DECK', '190121', '10', 'PR'],
                            ['DECK', '232909', '200', 'KG'],
                            ['DECK', '331130', '2', 'PKT'],
                            ['DECK', '331141', '12', 'PC'],
                            ['DECK', '450566', '4', 'TIN'],
                            ['DECK', '450702', '5', 'TIN'],
                            ['DECK', '470690', '24', 'PC'],
                            ['DECK', '470691', '24', 'PC'],
                            ['DECK', '470694', '24', 'PC'],
                            ['DECK', '471281', '6', 'RL'],
                            ['DECK', '471283', '6', 'RL'],
                            ['DECK', '510101', '24', 'PC'],
                            ['DECK', '510156', '24', 'PC'],
                            ['DECK', '510165', '24', 'PC'],
                            ['DECK', '510166', '24', 'PC'],
                            ['DECK', '510185', '5', 'PC'],
                            ['DECK', '510336', '24', 'PC'],
                            ['DECK', '510338', '24', 'PC'],
                            ['DECK', '510461', '30', 'SET'],
                            ['DECK', '510462', '100', 'PC'],
                            ['DECK', '510613', '10', 'PC'],
                            ['DECK', '510785', '20', 'PC'],
                            ['DECK', '510801', '18', 'PC'],
                            ['DECK', '510484', '6', 'PC'],
                            ['DECK', '510488', '10', 'PC'],
                            ['DECK', '510660', '10', 'PC'],
                            ['DECK', '510662', '10', 'PC'],
                            ['DECK', '510803', '30', 'PC'],
                            ['DECK', '550133', '4', 'BAG'],
                            ['DECK', '550920', '1', 'BOX'],
                            ['DECK', '550942', '4', 'BAG'],
                            ['DECK', '590468', '5', 'BOX'],
                            ['DECK', '611855', '2', 'PC'],
                            ['DECK', '614611', '300', 'SHT'],
                            ['DECK', '812725', '4', 'SET'],
                            ['DECK', '174124', '5', 'PC'],
                            ['DECK', '174141', '8', 'PC'],
                            ['DECK', '174143', '2', 'PC'],
                            ['DECK', '174153', '5', 'PC'],
                            ['DECK', '174280', '2', 'PC'],
                            ['DECK', '190107', '10', 'PR'],
                            ['DECK', '190631', '10', 'PC'],
                            ['DECK', '210108', '1', 'COIL'],
                            ['DECK', '210206', '1', 'COIL'],
                            ['DECK', '211207', '1', 'COIL'],
                            ['DECK', '211271', '2', 'COIL'],
                            ['DECK', '211302', '1', 'COIL'],
                            ['DECK', '211354', '2', 'COIL'],
                            ['DECK', '211452', '10', 'KG'],
                            ['DECK', '212251', '1', 'COIL'],
                            ['DECK', '212253', '1', 'COIL'],
                            ['DECK', '212257', '1', 'COIL'],
                            ['DECK', '212260', '1', 'COIL'],
                            ['DECK', '230157', '10', 'PC'],
                            ['DECK', '230158', '10', 'PC'],
                            ['DECK', '230162', '10', 'PC'],
                            ['DECK', '230323', '10', 'PC'],
                            ['DECK', '230324', '10', 'PC'],
                            ['DECK', '230325', '10', 'PC'],
                            ['DECK', '230519', '10', 'PC'],
                            ['DECK', '230520', '10', 'PC'],
                            ['DECK', '230521', '10', 'PC'],
                            ['DECK', '230820', '5', 'PC'],
                            ['DECK', '230821', '5', 'PC'],
                            ['DECK', '230822', '5', 'PC'],
                            ['DECK', '230823', '5', 'PC'],
                            ['DECK', '230824', '5', 'PC'],
                            ['DECK', '230948', '2', 'PC'],
                            ['DECK', '231602', '2', 'SET'],
                            ['DECK', '231652', '2', 'SET'],
                            ['DECK', '232012', '2', 'PC'],
                            ['DECK', '232111', '1', 'SET'],
                            ['DECK', '232121', '1', 'SET'],
                            ['DECK', '232152', '2', 'PC'],
                            ['DECK', '232157', '1', 'PC'],
                            ['DECK', '232161', '2', 'SHT'],
                            ['DECK', '232261', '2', 'PC'],
                            ['DECK', '232268', '5', 'PC'],
                            ['DECK', '232269', '5', 'PC'],
                            ['DECK', '232270', '5', 'PC'],
                            ['DECK', '232281', '1', 'SET'],
                            ['DECK', '232286', '50', 'PC'],
                            ['DECK', '232333', '1', 'PC'],
                            ['DECK', '232335', '1', 'PC'],
                            ['DECK', '232362', '12', 'PC'],
                            ['DECK', '232403', '2', 'SET'],
                            ['DECK', '232422', '30', 'PC'],
                            ['DECK', '232518', '5', 'BOX'],
                            ['DECK', '232705', '5', 'PC'],
                            ['DECK', '232706', '5', 'PC'],
                            ['DECK', '232754', '1', 'KG'],
                            ['DECK', '232758', '1', 'KG'],
                            ['DECK', '232760', '1', 'KG'],
                            ['DECK', '232762', '1', 'KG'],
                            ['DECK', '232946', '10', 'BAG'],
                            ['DECK', '232950', '6', 'BAG'],
                            ['DECK', '232956', '3', 'BAG'],
                            ['DECK', '233045', '4', 'SHT'],
                            ['DECK', '251402', '2', 'CAN'],
                            ['DECK', '270402', '1', 'PC'],
                            ['DECK', '270427', '2', 'COIL'],
                            ['DECK', '270456', '3', 'PC'],
                            ['DECK', '270581', '1', 'SET'],
                            ['DECK', '330131', '8', 'PC'],
                            ['DECK', '330143', '2', 'PC'],
                            ['DECK', '330189', '2', 'RL'],
                            ['DECK', '330252', '2', 'PC'],
                            ['DECK', '331101', '4', 'PC'],
                            ['DECK', '331171', '14', 'PC'],
                            ['DECK', '331176', '14', 'PC'],
                            ['DECK', '338695', '30', 'PC'],
                            ['DECK', '350102', '3', 'SET'],
                            ['DECK', '350181', '2', 'COIL'],
                            ['DECK', '350206', '2', 'COIL'],
                            ['DECK', '350331', '2', 'COIL'],
                            ['DECK', '350331', '2', 'SET'],
                            ['DECK', '351008', '1', 'PC'],
                            ['DECK', '352001', '2', 'COIL'],
                            ['DECK', '352001', '2', 'SET'],
                            ['DECK', '370213', '2', 'PC'],
                            ['DECK', '370251', '4', 'PC'],
                            ['DECK', '370307', '1', 'PC'],
                            ['DECK', '370343', '1', 'PC'],
                            ['DECK', '371002', '1', 'PC'],
                            ['DECK', '371007', '2', 'SET'],
                            ['DECK', '371013', '2', 'PC'],
                            ['DECK', '371016', '2', 'PC'],
                            ['DECK', '371021', '4', 'PC'],
                            ['DECK', '371028', '1', 'PC'],
                            ['DECK', '371031', '2', 'PC'],
                            ['DECK', '371581', '10', 'SET'],
                            ['DECK', '371587', '5', 'PC'],
                            ['DECK', '470245', '1', 'PC'],
                            ['DECK', '470702', '1', 'BOX'],
                            ['DECK', '471301', '2', 'PKT'],
                            ['DECK', '471302', '2', 'PKT'],
                            ['DECK', '471566', '5', 'PC'],
                            ['DECK', '471630', '2', 'PC'],
                            ['DECK', '490511', '10', 'PC'],
                            ['DECK', '490512', '10', 'PC'],
                            ['DECK', '490513', '20', 'PC'],
                            ['DECK', '510501', '4', 'PC'],
                            ['DECK', '511081', '1', 'RL'],
                            ['DECK', '511086', '1', 'RL'],
                            ['DECK', '550262', '2', 'CASE'],
                            ['DECK', '550532', '2', 'CAN'],
                            ['DECK', '232521', '10', 'TIN'],
                            ['DECK', '590463', '1', 'SET'],
                            ['DECK', '590467', '1', 'BOX'],
                            ['DECK', '590741', '1', 'SET'],
                            ['DECK', '591021', '1', 'SET'],
                            ['DECK', '591602', '1', 'PC'],
                            ['DECK', '610511', '1', 'SET'],
                            ['DECK', '610665', '1', 'PC'],
                            ['DECK', '611331', '1', 'PC'],
                            ['DECK', '611333', '1', 'PC'],
                            ['DECK', '611673', '2', 'PR'],
                            ['DECK', '611843', '1', 'PC'],
                            ['DECK', '611901', '1', 'SET'],
                            ['DECK', '611971', '1', 'SET'],
                            ['DECK', '612263', '2', 'PC'],
                            ['DECK', '612266', '2', 'PC'],
                            ['DECK', '612268', '2', 'PC'],
                            ['DECK', '612274', '2', 'PC'],
                            ['DECK', '612276', '2', 'PC'],
                            ['DECK', '612416', '1', 'SET'],
                            ['DECK', '612504', '3', 'PC'],
                            ['DECK', '612507', '2', 'PC'],
                            ['DECK', '612523', '1', 'PC'],
                            ['DECK', '612526', '1', 'PC'],
                            ['DECK', '612602', '3', 'PC'],
                            ['DECK', '612612', '10', 'PC'],
                            ['DECK', '612785', '10', 'PC'],
                            ['DECK', '612860', '1', 'PC'],
                            ['DECK', '612884', '1', 'PC'],
                            ['DECK', '612907', '2', 'PC'],
                            ['DECK', '613133', '2', 'SET'],
                            ['DECK', '613140', '2', 'SET'],
                            ['DECK', '613242', '10', 'PC'],
                            ['DECK', '613246', '10', 'PC'],
                            ['DECK', '613404', '2', 'PC'],
                            ['DECK', '613426', '2', 'PC'],
                            ['DECK', '613433', '2', 'DOZ'],
                            ['DECK', '613653', '1', 'PC'],
                            ['DECK', '613686', '10', 'PC'],
                            ['DECK', '613689', '4', 'PC'],
                            ['DECK', '614007', '1', 'PC'],
                            ['DECK', '614010', '5', 'PC'],
                            ['DECK', '614053', '25', 'PC'],
                            ['DECK', '614056', '25', 'PC'],
                            ['DECK', '614057', '12', 'PC'],
                            ['DECK', '615007', '1', 'SET'],
                            ['DECK', '615009', '1', 'SET'],
                            ['DECK', '616533', '1', 'PC'],
                            ['DECK', '616535', '1', 'PC'],
                            ['DECK', '617104', '1', 'SET'],
                            ['DECK', '617123', '1', 'SET'],
                            ['DECK', '617137', '1', 'SET'],
                            ['DECK', '617152', '1', 'SET'],
                            ['DECK', '617191', '1', 'PC'],
                            ['DECK', '617516', '2', 'PC'],
                            ['DECK', '617624', '20', 'PC'],
                            ['DECK', '617626', '10', 'PC'],
                            ['DECK', '617680', '2', 'PC'],
                            ['DECK', '617721', '5', 'PC'],
                            ['DECK', '650503', '1', 'PC'],
                            ['DECK', '650822', '1', 'PC'],
                            ['DECK', '650823', '1', 'PC'],
                            ['DECK', '650839', '1', 'PC'],
                            ['DECK', '650878', '2', 'PC'],
                            ['DECK', '650890', '20', 'PC'],
                            ['DECK', '651342', '2', 'PC'],
                            ['DECK', '651371', '1', 'SET'],
                            ['DECK', '651701', '5', 'PC'],
                            ['DECK', '651725', '10', 'SET'],
                            ['DECK', '671126', '10', 'KG'],
                            ['DECK', '696604', '1', 'PKT'],
                            ['DECK', '696624', '1', 'PKT'],
                            ['DECK', '795433', '5', 'PC'],
                            ['DECK', '795436', '5', 'PC'],
                            ['DECK', '812702', '2', 'PC']
            ]
    qtydup_table_eng = [
                            ['dept', 'impa', 'qty', 'unit'],
                            ['ENGINE', '174175', '30', 'BAG'],
                            ['ENGINE', '174176', '6', 'BAG'],
                            ['ENGINE', '174275', '4', 'PC'],
                            ['ENGINE', '174277', '24', 'PC'],
                            ['ENGINE', '174280', '3', 'PC'],
                            ['ENGINE', '190101', '50', 'DOZ'],
                            ['ENGINE', '190122', '10', 'PR'],
                            ['ENGINE', '190132', '10', 'PR'],
                            ['ENGINE', '232422', '35', 'PC'],
                            ['ENGINE', '232907', '100', 'KG'],
                            ['ENGINE', '232909', '200', 'KG'],
                            ['ENGINE', '331130', '2', 'PKT'],
                            ['ENGINE', '331141', '10', 'PR'],
                            ['ENGINE', '331156', '27', 'PR'],
                            ['ENGINE', '331157', '20', 'PC'],
                            ['ENGINE', '450107', '10', 'CAN'],
                            ['ENGINE', '450514', '5', 'CAN'],
                            ['ENGINE', '450566', '4', 'TIN'],
                            ['ENGINE', '450702', '8', 'TIN'],
                            ['ENGINE', '470603', '30', 'PC'],
                            ['ENGINE', '470701', '1', 'BOX'],
                            ['ENGINE', '470702', '1', 'BOX'],
                            ['ENGINE', '470882', '3', 'PC'],
                            ['ENGINE', '471283', '6', 'RL'],
                            ['ENGINE', '471291', '5', 'PC'],
                            ['ENGINE', '471292', '5', 'PC'],
                            ['ENGINE', '471293', '5', 'PC'],
                            ['ENGINE', '471294', '3', 'PC'],
                            ['ENGINE', '510167', '10', 'PC'],
                            ['ENGINE', '510461', '3', 'SET'],
                            ['ENGINE', '510462', '24', 'PC'],
                            ['ENGINE', '510623', '6', 'PC'],
                            ['ENGINE', '550133', '4', 'BAG'],
                            ['ENGINE', '550271', '3', 'TIN'],
                            ['ENGINE', '795433', '5', 'PC'],
                            ['ENGINE', '795433', '5', 'PC'],
                            ['ENGINE', '795433', '5', 'PC'],
                            ['ENGINE', '795501', '4', 'TIN'],
                            ['ENGINE', '812251', '3', 'SET'],
                            ['ENGINE', '812501', '20', 'PC'],
                            ['ENGINE', '812602', '8', 'TUB'],
                            ['ENGINE', '174012', '1', 'SET'],
                            ['ENGINE', '174029', '1', 'PC'],
                            ['ENGINE', '174124', '5', 'PC'],
                            ['ENGINE', '174276', '10', 'PC'],
                            ['ENGINE', '190107', '9', 'PC'],
                            ['ENGINE', '190413', '9', 'PC'],
                            ['ENGINE', '190631', '9', 'PC'],
                            ['ENGINE', '211204', '1', 'COIL'],
                            ['ENGINE', '211205', '1', 'COIL'],
                            ['ENGINE', '211402', '10', 'HNK'],
                            ['ENGINE', '211453', '10', 'KG'],
                            ['ENGINE', '230155', '4', 'PC'],
                            ['ENGINE', '230157', '4', 'PC'],
                            ['ENGINE', '230159', '4', 'PC'],
                            ['ENGINE', '230343', '4', 'PC'],
                            ['ENGINE', '230344', '4', 'PC'],
                            ['ENGINE', '230345', '4', 'PC'],
                            ['ENGINE', '230346', '4', 'PC'],
                            ['ENGINE', '232604', '2', 'PC'],
                            ['ENGINE', '232705', '10', 'PC'],
                            ['ENGINE', '232756', '1', 'KG'],
                            ['ENGINE', '232760', '1', 'KG'],
                            ['ENGINE', '232946', '10', 'BAG'],
                            ['ENGINE', '232978', '7', 'BAG'],
                            ['ENGINE', '270613', '3', 'PC'],
                            ['ENGINE', '310103', '2', 'PC'],
                            ['ENGINE', '331101', '2', 'PC'],
                            ['ENGINE', '331111', '5', 'PC'],
                            ['ENGINE', '331151', '9', 'PC'],
                            ['ENGINE', '331176', '4', 'PC'],
                            ['ENGINE', '350102', '50', 'MTR'],
                            ['ENGINE', '350117', '1', 'COIL'],
                            ['ENGINE', '350122', '1', 'COIL'],
                            ['ENGINE', '350237', '1', 'COIL'],
                            ['ENGINE', '350252', '15', 'MTR'],
                            ['ENGINE', '450511', '2', 'CAN'],
                            ['ENGINE', '450562', '1', 'CAN'],
                            ['ENGINE', '470245', '1', 'PC'],
                            ['ENGINE', '470544', '1', 'PC'],
                            ['ENGINE', '470678', '1', 'PKT'],
                            ['ENGINE', '471301', '1', 'BDL'],
                            ['ENGINE', '471302', '1', 'BDL'],
                            ['ENGINE', '471567', '2', 'PC'],
                            ['ENGINE', '471651', '6', 'PC'],
                            ['ENGINE', '471652', '6', 'PC'],
                            ['ENGINE', '471653', '6', 'PC'],
                            ['ENGINE', '471654', '6', 'PC'],
                            ['ENGINE', '490503', '5', 'PC'],
                            ['ENGINE', '490506', '5', 'PC'],
                            ['ENGINE', '510165', '12', 'PC'],
                            ['ENGINE', '510185', '4', 'PC'],
                            ['ENGINE', '510231', '3', 'PC'],
                            ['ENGINE', '510319', '3', 'PC'],
                            ['ENGINE', '510660', '5', 'PC'],
                            ['ENGINE', '510662', '5', 'PC'],
                            ['ENGINE', '510785', '5', 'PC'],
                            ['ENGINE', '510807', '8', 'PC'],
                            ['ENGINE', '550101', '10', 'PC'],
                            ['ENGINE', '550352', '5', 'CAN'],
                            ['ENGINE', '550358', '5', 'CAN'],
                            ['ENGINE', '550689', '3', 'BTL'],
                            ['ENGINE', '550847', '2', 'PC'],
                            ['ENGINE', '550951', '3', 'CAN'],
                            ['ENGINE', '590105', '1', 'SET'],
                            ['ENGINE', '590229', '1', 'PC'],
                            ['ENGINE', '590230', '1', 'PC'],
                            ['ENGINE', '590231', '1', 'PC'],
                            ['ENGINE', '590233', '1', 'PC'],
                            ['ENGINE', '590235', '1', 'PC'],
                            ['ENGINE', '590237', '1', 'PC'],
                            ['ENGINE', '590238', '1', 'PC'],
                            ['ENGINE', '590239', '1', 'PC'],
                            ['ENGINE', '590461', '1', 'PC'],
                            ['ENGINE', '590721', '1', 'SET'],
                            ['ENGINE', '591003', '1', 'SET'],
                            ['ENGINE', '591021', '1', 'SET'],
                            ['ENGINE', '591152', '1', 'SET'],
                            ['ENGINE', '591402', '1', 'SET'],
                            ['ENGINE', '591481', '2', 'SET'],
                            ['ENGINE', '610120', '1', 'SET'],
                            ['ENGINE', '610153', '1', 'SET'],
                            ['ENGINE', '610252', '1', 'PC'],
                            ['ENGINE', '610255', '1', 'PC'],
                            ['ENGINE', '610257', '1', 'PC'],
                            ['ENGINE', '610262', '1', 'PC'],
                            ['ENGINE', '610279', '1', 'PC'],
                            ['ENGINE', '610282', '1', 'PC'],
                            ['ENGINE', '610284', '1', 'PC'],
                            ['ENGINE', '610289', '1', 'PC'],
                            ['ENGINE', '610405', '1', 'PC'],
                            ['ENGINE', '610552', '2', 'PC'],
                            ['ENGINE', '610556', '2', 'PC'],
                            ['ENGINE', '610565', '2', 'PC'],
                            ['ENGINE', '610566', '2', 'PC'],
                            ['ENGINE', '610568', '2', 'PC'],
                            ['ENGINE', '610570', '2', 'PC'],
                            ['ENGINE', '610571', '2', 'PC'],
                            ['ENGINE', '610575', '2', 'PC'],
                            ['ENGINE', '610581', '2', 'PC'],
                            ['ENGINE', '610583', '2', 'PC'],
                            ['ENGINE', '610585', '2', 'PC'],
                            ['ENGINE', '610588', '1', 'PC'],
                            ['ENGINE', '610602', '1', 'PC'],
                            ['ENGINE', '610606', '2', 'PC'],
                            ['ENGINE', '610609', '2', 'PC'],
                            ['ENGINE', '610611', '2', 'PC'],
                            ['ENGINE', '610612', '2', 'PC'],
                            ['ENGINE', '610614', '2', 'PC'],
                            ['ENGINE', '610616', '2', 'PC'],
                            ['ENGINE', '610618', '2', 'PC'],
                            ['ENGINE', '610620', '2', 'PC'],
                            ['ENGINE', '610621', '2', 'PC'],
                            ['ENGINE', '610623', '2', 'PC'],
                            ['ENGINE', '610651', '2', 'PC'],
                            ['ENGINE', '610654', '2', 'PC'],
                            ['ENGINE', '610656', '2', 'PC'],
                            ['ENGINE', '610657', '2', 'PC'],
                            ['ENGINE', '610659', '2', 'PC'],
                            ['ENGINE', '610661', '2', 'PC'],
                            ['ENGINE', '610663', '2', 'PC'],
                            ['ENGINE', '610665', '2', 'PC'],
                            ['ENGINE', '610666', '2', 'PC'],
                            ['ENGINE', '610668', '2', 'PC'],
                            ['ENGINE', '610702', '1', 'PC'],
                            ['ENGINE', '610707', '1', 'PC'],
                            ['ENGINE', '610708', '2', 'PC'],
                            ['ENGINE', '610712', '2', 'PC'],
                            ['ENGINE', '610715', '2', 'PC'],
                            ['ENGINE', '610720', '2', 'PC'],
                            ['ENGINE', '610860', '1', 'PC'],
                            ['ENGINE', '610861', '1', 'PC'],
                            ['ENGINE', '610863', '1', 'PC'],
                            ['ENGINE', '610865', '1', 'PC'],
                            ['ENGINE', '610869', '1', 'PC'],
                            ['ENGINE', '610870', '1', 'PC'],
                            ['ENGINE', '610872', '2', 'PC'],
                            ['ENGINE', '611060', '1', 'PC'],
                            ['ENGINE', '611066', '1', 'PC'],
                            ['ENGINE', '611281', '1', 'SET'],
                            ['ENGINE', '611305', '1', 'PC'],
                            ['ENGINE', '611331', '1', 'PC'],
                            ['ENGINE', '611332', '1', 'PC'],
                            ['ENGINE', '611333', '1', 'PC'],
                            ['ENGINE', '611334', '1', 'PC'],
                            ['ENGINE', '611335', '1', 'PC'],
                            ['ENGINE', '611601', '1', 'PC'],
                            ['ENGINE', '611602', '1', 'PC'],
                            ['ENGINE', '611651', '1', 'PC'],
                            ['ENGINE', '611653', '1', 'PC'],
                            ['ENGINE', '611654', '1', 'PC'],
                            ['ENGINE', '611657', '1', 'PC'],
                            ['ENGINE', '611694', '1', 'PC'],
                            ['ENGINE', '611711', '1', 'PC'],
                            ['ENGINE', '611774', '1', 'PC'],
                            ['ENGINE', '611783', '1', 'PC'],
                            ['ENGINE', '611841', '1', 'PC'],
                            ['ENGINE', '611843', '1', 'PC'],
                            ['ENGINE', '611855', '1', 'SET'],
                            ['ENGINE', '612001', '1', 'PC'],
                            ['ENGINE', '612004', '1', 'SET'],
                            ['ENGINE', '612030', '1', 'SET'],
                            ['ENGINE', '612287', '1', 'PC'],
                            ['ENGINE', '612289', '1', 'PC'],
                            ['ENGINE', '612291', '1', 'PC'],
                            ['ENGINE', '612296', '1', 'PC'],
                            ['ENGINE', '612298', '1', 'PC'],
                            ['ENGINE', '612324', '1', 'PC'],
                            ['ENGINE', '612325', '1', 'PC'],
                            ['ENGINE', '612417', '1', 'SET'],
                            ['ENGINE', '612504', '2', 'PC'],
                            ['ENGINE', '612508', '2', 'PC'],
                            ['ENGINE', '612524', '1', 'PC'],
                            ['ENGINE', '612527', '1', 'PC'],
                            ['ENGINE', '612531', '2', 'PC'],
                            ['ENGINE', '612602', '2', 'PC'],
                            ['ENGINE', '612612', '2', 'PC'],
                            ['ENGINE', '612617', '1', 'PC'],
                            ['ENGINE', '612644', '1', 'PC'],
                            ['ENGINE', '612684', '1', 'PC'],
                            ['ENGINE', '612724', '2', 'PC'],
                            ['ENGINE', '612763', '1', 'PC'],
                            ['ENGINE', '612860', '1', 'PC'],
                            ['ENGINE', '612882', '1', 'PC'],
                            ['ENGINE', '612904', '1', 'PC'],
                            ['ENGINE', '612907', '1', 'PC'],
                            ['ENGINE', '612913', '1', 'PC'],
                            ['ENGINE', '612932', '1', 'PC'],
                            ['ENGINE', '612933', '1', 'PC'],
                            ['ENGINE', '613008', '2', 'PC'],
                            ['ENGINE', '613017', '1', 'PC'],
                            ['ENGINE', '613018', '1', 'PC'],
                            ['ENGINE', '613028', '1', 'PC'],
                            ['ENGINE', '613056', '1', 'SET'],
                            ['ENGINE', '613061', '1', 'PC'],
                            ['ENGINE', '613106', '1', 'SET'],
                            ['ENGINE', '613121', '1', 'SET'],
                            ['ENGINE', '613186', '1', 'PC'],
                            ['ENGINE', '613202', '2', 'PC'],
                            ['ENGINE', '613232', '1', 'PC'],
                            ['ENGINE', '613242', '3', 'PC'],
                            ['ENGINE', '613265', '5', 'PC'],
                            ['ENGINE', '613281', '1', 'PC'],
                            ['ENGINE', '613426', '2', 'SET'],
                            ['ENGINE', '613433', '3', 'DOZ'],
                            ['ENGINE', '613774', '1', 'PC'],
                            ['ENGINE', '614007', '1', 'SET'],
                            ['ENGINE', '614010', '3', 'PC'],
                            ['ENGINE', '614011', '1', 'PC'],
                            ['ENGINE', '614025', '1', 'SET'],
                            ['ENGINE', '614031', '1', 'PC'],
                            ['ENGINE', '614032', '1', 'PC'],
                            ['ENGINE', '614202', '2', 'CAN'],
                            ['ENGINE', '614192', '3', 'CAN'],
                            ['ENGINE', '614193', '3', 'CAN'],
                            ['ENGINE', '614215', '3', 'CAN'],
                            ['ENGINE', '614311', '2', 'PC'],
                            ['ENGINE', '614338', '2', 'PC'],
                            ['ENGINE', '614366', '2', 'PC'],
                            ['ENGINE', '614394', '2', 'PC'],
                            ['ENGINE', '614425', '2', 'PC'],
                            ['ENGINE', '614534', '1', 'SET'],
                            ['ENGINE', '614562', '3', 'PC'],
                            ['ENGINE', '614611', '100', 'SHT'],
                            ['ENGINE', '614684', '50', 'SHT'],
                            ['ENGINE', '614688', '75', 'SHT'],
                            ['ENGINE', '614692', '75', 'SHT'],
                            ['ENGINE', '614695', '30', 'SHT'],
                            ['ENGINE', '614803', '10', 'PC'],
                            ['ENGINE', '614878', '5', 'PC'],
                            ['ENGINE', '615006', '2', 'PC'],
                            ['ENGINE', '615007', '1', 'PC'],
                            ['ENGINE', '615009', '1', 'PC'],
                            ['ENGINE', '615051', '1', 'SET'],
                            ['ENGINE', '615052', '1', 'SET'],
                            ['ENGINE', '615053', '1', 'SET'],
                            ['ENGINE', '615113', '1', 'SET'],
                            ['ENGINE', '615115', '1', 'SET'],
                            ['ENGINE', '616530', '1', 'PC'],
                            ['ENGINE', '616532', '1', 'PC'],
                            ['ENGINE', '616534', '1', 'PC'],
                            ['ENGINE', '616601', '1', 'SET'],
                            ['ENGINE', '616673', '1', 'SET'],
                            ['ENGINE', '617016', '1', 'SET'],
                            ['ENGINE', '617017', '3', 'PC'],
                            ['ENGINE', '617020', '1', 'SET'],
                            ['ENGINE', '617021', '5', 'PC'],
                            ['ENGINE', '617026', '1', 'SET'],
                            ['ENGINE', '617103', '1', 'PC'],
                            ['ENGINE', '617137', '1', 'PC'],
                            ['ENGINE', '617707', '2', 'SET'],
                            ['ENGINE', '617624', '20', 'PC'],
                            ['ENGINE', '617680', '3', 'PC'],
                            ['ENGINE', '617718', '1', 'PC'],
                            ['ENGINE', '617721', '5', 'PC'],
                            ['ENGINE', '617745', '2', 'PC'],
                            ['ENGINE', '617746', '2', 'PC'],
                            ['ENGINE', '617747', '2', 'PC'],
                            ['ENGINE', '630118', '2', 'PC'],
                            ['ENGINE', '630128', '2', 'PC'],
                            ['ENGINE', '630138', '2', 'PC'],
                            ['ENGINE', '630148', '2', 'PC'],
                            ['ENGINE', '630158', '2', 'PC'],
                            ['ENGINE', '630168', '2', 'PC'],
                            ['ENGINE', '630178', '2', 'PC'],
                            ['ENGINE', '630198', '2', 'PC'],
                            ['ENGINE', '630201', '2', 'PC'],
                            ['ENGINE', '630203', '2', 'PC'],
                            ['ENGINE', '630307', '2', 'PC'],
                            ['ENGINE', '630309', '1', 'PC'],
                            ['ENGINE', '630311', '1', 'PC'],
                            ['ENGINE', '630313', '2', 'PC'],
                            ['ENGINE', '630319', '2', 'PC'],
                            ['ENGINE', '630323', '1', 'PC'],
                            ['ENGINE', '630331', '1', 'PC'],
                            ['ENGINE', '631018', '1', 'SET'],
                            ['ENGINE', '631020', '1', 'SET'],
                            ['ENGINE', '631022', '1', 'SET'],
                            ['ENGINE', '631024', '1', 'SET'],
                            ['ENGINE', '632203', '1', 'SET'],
                            ['ENGINE', '650301', '1', 'PC'],
                            ['ENGINE', '650512', '1', 'PC'],
                            ['ENGINE', '650553', '1', 'SET'],
                            ['ENGINE', '650601', '1', 'PC'],
                            ['ENGINE', '650602', '1', 'PC'],
                            ['ENGINE', '650804', '1', 'PC'],
                            ['ENGINE', '650823', '2', 'PC'],
                            ['ENGINE', '650838', '1', 'PC'],
                            ['ENGINE', '650873', '4', 'PC'],
                            ['ENGINE', '650890', '4', 'PC'],
                            ['ENGINE', '651021', '1', 'SET'],
                            ['ENGINE', '651701', '1', 'PC'],
                            ['ENGINE', '651723', '2', 'PC'],
                            ['ENGINE', '651725', '2', 'PC'],
                            ['ENGINE', '670253', '4', 'PC'],
                            ['ENGINE', '670254', '4', 'PC'],
                            ['ENGINE', '670256', '4', 'PC'],
                            ['ENGINE', '670257', '4', 'PC'],
                            ['ENGINE', '670404', '4', 'PC'],
                            ['ENGINE', '670410', '4', 'PC'],
                            ['ENGINE', '670417', '4', 'PC'],
                            ['ENGINE', '670602', '5', 'PC'],
                            ['ENGINE', '670603', '5', 'PC'],
                            ['ENGINE', '670607', '5', 'PC'],
                            ['ENGINE', '670608', '5', 'PC'],
                            ['ENGINE', '670611', '4', 'PC'],
                            ['ENGINE', '670705', '2', 'PC'],
                            ['ENGINE', '670707', '2', 'PC'],
                            ['ENGINE', '670708', '2', 'PC'],
                            ['ENGINE', '670709', '2', 'PC'],
                            ['ENGINE', '670710', '2', 'PC'],
                            ['ENGINE', '670856', '2', 'SHT'],
                            ['ENGINE', '670784', '2', 'SHT'],
                            ['ENGINE', '671118', '2', 'KG'],
                            ['ENGINE', '671861', '2', 'KG'],
                            ['ENGINE', '671901', '1', 'SHT'],
                            ['ENGINE', '671906', '1', 'SHT'],
                            ['ENGINE', '671909', '1', 'SHT'],
                            ['ENGINE', '692005', '5', 'PC'],
                            ['ENGINE', '692006', '5', 'PC'],
                            ['ENGINE', '692007', '5', 'PC'],
                            ['ENGINE', '692008', '5', 'PC'],
                            ['ENGINE', '692009', '5', 'PC'],
                            ['ENGINE', '692030', '5', 'PC'],
                            ['ENGINE', '692031', '5', 'PC'],
                            ['ENGINE', '696702', '10', 'PC'],
                            ['ENGINE', '696703', '10', 'PC'],
                            ['ENGINE', '696704', '10', 'PC'],
                            ['ENGINE', '696705', '10', 'PC'],
                            ['ENGINE', '696706', '10', 'PC'],
                            ['ENGINE', '696707', '10', 'PC'],
                            ['ENGINE', '696708', '10', 'PC'],
                            ['ENGINE', '696710', '10', 'PC'],
                            ['ENGINE', '710104', '2', 'PC'],
                            ['ENGINE', '710105', '2', 'PC'],
                            ['ENGINE', '710106', '2', 'PC'],
                            ['ENGINE', '710107', '2', 'PC'],
                            ['ENGINE', '710108', '2', 'PC'],
                            ['ENGINE', '710109', '2', 'PC'],
                            ['ENGINE', '710110', '2', 'PC'],
                            ['ENGINE', '710111', '2', 'PC'],
                            ['ENGINE', '710113', '2', 'PC'],
                            ['ENGINE', '711502', '1', 'PC'],
                            ['ENGINE', '711504', '1', 'PC'],
                            ['ENGINE', '711506', '1', 'PC'],
                            ['ENGINE', '711507', '1', 'PC'],
                            ['ENGINE', '711509', '1', 'PC'],
                            ['ENGINE', '731121', '5', 'PC'],
                            ['ENGINE', '731122', '5', 'PC'],
                            ['ENGINE', '731123', '5', 'PC'],
                            ['ENGINE', '731124', '5', 'PC'],
                            ['ENGINE', '731125', '5', 'PC'],
                            ['ENGINE', '731126', '5', 'PC'],
                            ['ENGINE', '731127', '5', 'PC'],
                            ['ENGINE', '731128', '5', 'PC'],
                            ['ENGINE', '731130', '5', 'PC'],
                            ['ENGINE', '734003', '10', 'PC'],
                            ['ENGINE', '734005', '10', 'PC'],
                            ['ENGINE', '734007', '10', 'PC'],
                            ['ENGINE', '734009', '10', 'PC'],
                            ['ENGINE', '734022', '10', 'PC'],
                            ['ENGINE', '734023', '10', 'PC'],
                            ['ENGINE', '734024', '10', 'PC'],
                            ['ENGINE', '734062', '10', 'PC'],
                            ['ENGINE', '734063', '10', 'PC'],
                            ['ENGINE', '734064', '10', 'PC'],
                            ['ENGINE', '734065', '10', 'PC'],
                            ['ENGINE', '734602', '6', 'PC'],
                            ['ENGINE', '734603', '6', 'PC'],
                            ['ENGINE', '734604', '6', 'PC'],
                            ['ENGINE', '734605', '6', 'PC'],
                            ['ENGINE', '734606', '6', 'PC'],
                            ['ENGINE', '734608', '6', 'PC'],
                            ['ENGINE', '734609', '6', 'PC'],
                            ['ENGINE', '734610', '6', 'PC'],
                            ['ENGINE', '734611', '6', 'PC'],
                            ['ENGINE', '791502', '1', 'CASE'],
                            ['ENGINE', '792166', '3', 'SET'],
                            ['ENGINE', '792201', '4', 'PC'],
                            ['ENGINE', '792202', '5', 'PC'],
                            ['ENGINE', '792251', '25', 'PC'],
                            ['ENGINE', '792661', '1', 'PC'],
                            ['ENGINE', '792806', '6', 'PC'],
                            ['ENGINE', '792961', '6', 'PC'],
                            ['ENGINE', '793737', '1', 'PC'],
                            ['ENGINE', '794122', '50', 'MTR'],
                            ['ENGINE', '794132', '50', 'MTR'],
                            ['ENGINE', '794203', '50', 'MTR'],
                            ['ENGINE', '794391', '1', 'SET'],
                            ['ENGINE', '794761', '2', 'SET'],
                            ['ENGINE', '795013', '1', 'PC'],
                            ['ENGINE', '795016', '1', 'PC'],
                            ['ENGINE', '795170', '1', 'COIL'],
                            ['ENGINE', '795172', '5', 'BAR'],
                            ['ENGINE', '795174', '1', 'PC'],
                            ['ENGINE', '795532', '2', 'PR'],
                            ['ENGINE', '795741', '1', 'SET'],
                            ['ENGINE', '811111', '2', 'SHT'],
                            ['ENGINE', '811112', '2', 'SHT'],
                            ['ENGINE', '811113', '1', 'SHT'],
                            ['ENGINE', '811114', '1', 'SHT'],
                            ['ENGINE', '811126', '1', 'SHT'],
                            ['ENGINE', '811128', '1', 'SHT'],
                            ['ENGINE', '811502', '1', 'COIL'],
                            ['ENGINE', '812522', '1', 'LGH'],
                            ['ENGINE', '812527', '1', 'LGH'],
                            ['ENGINE', '812531', '1', 'SHT'],
                            ['ENGINE', '812702', '1', 'PC'],
                            ['ENGINE', '813070', '1', 'SET'],
                            ['ENGINE', '813072', '1', 'SET'],
                            ['ENGINE', '813080', '1', 'SET'],
                            ['ENGINE', '813102', '20', 'MTR'],
                            ['ENGINE', '850280', '1', 'SET'],
                            ['ENGINE', '850286', '3', 'PC'],
                            ['ENGINE', '850288', '3', 'BOX'],
                            ['ENGINE', '850403', '1', 'KG'],
                            ['ENGINE', '850405', '1', 'KG'],
                            ['ENGINE', '850453', '1', 'KG'],
                            ['ENGINE', '850652', '2', 'CASE'],
                            ['ENGINE', '850663', '2', 'CASE'],
                            ['ENGINE', '850682', '2', 'CASE'],
                            ['ENGINE', '851032', '1', 'SET'],
                            ['ENGINE', '851121', '2', 'PC'],
                            ['ENGINE', '851143', '10', 'SHT'],
                            ['ENGINE', '851146', '10', 'SHT'],
                            ['ENGINE', '851163', '2', 'PR'],
                            ['ENGINE', '851166', '1', 'PR'],
                            ['ENGINE', '851167', '2', 'PC'],
                            ['ENGINE', '851168', '2', 'PR'],
                            ['ENGINE', '792401', '100', 'PC'],
                            ['ENGINE', '792402', '100', 'PC'],
                            ['ENGINE', '792403', '100', 'PC'],
                            ['ENGINE', '792410', '50', 'PC'],
                            ['ENGINE', '795431', '20', 'PC'],
                            ['ENGINE', '611732', '1', 'SET'],
                            ['ENGINE', '790403', '10', 'PC'],
                            ['ENGINE', '790404', '20', 'PC'],
                            ['ENGINE', '790105', '30', 'PC'],
                            ['ENGINE', '790106', '50', 'PC'],
                            ['ENGINE', '790108', '25', 'PC'],
                            ['ENGINE', '790184', '30', 'PC'],
                            ['ENGINE', '790185', '30', 'PC'],
                            ['ENGINE', '790201', '6', 'PC'],
                            ['ENGINE', '790232', '10', 'PC'],
                            ['ENGINE', '790706', '30', 'PC'],
                            ['ENGINE', '790742', '10', 'PC'],
                            ['ENGINE', '790924', '6', 'PC'],
                            ['ENGINE', '790925', '6', 'PC'],
                            ['ENGINE', '791115', '4', 'PC'],
                            ['ENGINE', '791117', '4', 'PC'],
                            ['ENGINE', '791123', '3', 'PC'],
                            ['ENGINE', '791403', '20', 'PC'],
                            ['ENGINE', '791405', '30', 'PC'],
                            ['ENGINE', '791407', '100', 'PC'],
                            ['ENGINE', '791418', '50', 'PC'],
                            ['ENGINE', '791503', '25', 'PC'],
                            ['ENGINE', '791504', '50', 'PC'],
                            ['ENGINE', '791505', '25', 'PC'],
                            ['ENGINE', '791803', '4', 'PC'],
                            ['ENGINE', '792248', '8', 'PC'],
                            ['ENGINE', '792254', '5', 'PC'],
                            ['ENGINE', '792255', '20', 'PC'],
                            ['ENGINE', '792801', '4', 'PC'],
                            ['ENGINE', '792802', '4', 'PC'],
                            ['ENGINE', '792808', '2', 'PC'],
                            ['ENGINE', '792836', '2', 'SET'],
                            ['ENGINE', '792977', '2', 'PC'],
                            ['ENGINE', '793880', '1', 'PC'],
                            ['ENGINE', '794392', '1', 'SET'],
                            ['ENGINE', '794393', '2', 'SET'],
                            ['ENGINE', '794637', '1', 'PC'],
                            ['ENGINE', '794851', '100', 'PC'],
                            ['ENGINE', '794852', '100', 'PC'],
                            ['ENGINE', '794853', '100', 'PC'],
                            ['ENGINE', '794881', '6', 'PC'],
                            ['ENGINE', '794882', '6', 'PC'],
                            ['ENGINE', '795506', '3', 'TIN'],
                            ['ENGINE', '795730', '1', 'PC'],
                            ['ENGINE', '795776', '1', 'SET']
            ]
    qtydup_table_sta = [
                            ['dept', 'impa', 'qty', 'unit'],
                            ['STATIONERY', '470110', '20', 'VOL'],
                            ['STATIONERY', '470127', '10', 'PC'],
                            ['STATIONERY', '470138', '20', 'VOL'],
                            ['STATIONERY', '470426', '12', 'PC'],
                            ['STATIONERY', '470457', '1', 'PKT'],
                            ['STATIONERY', '470507', '12', 'PC'],
                            ['STATIONERY', '470523', '12', 'PC'],
                            ['STATIONERY', '470603', '30', 'PC'],
                            ['STATIONERY', '470606', '12', 'PC'],
                            ['STATIONERY', '470607', '6', 'PC'],
                            ['STATIONERY', '470865', '20', 'PC'],
                            ['STATIONERY', '470871', '3', 'PC'],
                            ['STATIONERY', '470881', '6', 'RL'],
                            ['STATIONERY', '471082', '0', 'BTL'],
                            ['STATIONERY', '471092', '6', 'PC'],
                            ['STATIONERY', '471101', '10', 'BOX'],
                            ['STATIONERY', '471252', '6', 'PC'],
                            ['STATIONERY', '471268', '5', 'PC'],
                            ['STATIONERY', '471269', '5', 'PC'],
                            ['STATIONERY', '472186', '30', 'PKT'],
                            ['STATIONERY', '370854', '4', 'VOL'],
                            ['STATIONERY', '370857', '2', 'VOL'],
                            ['STATIONERY', '370861', '3', 'VOL'],
                            ['STATIONERY', '370862', '4', 'VOL'],
                            ['STATIONERY', '370871', '2', 'VOL'],
                            ['STATIONERY', '370874', '3', 'VOL'],
                            ['STATIONERY', '370876', '3', 'VOL'],
                            ['STATIONERY', '370877', '2', 'VOL'],
                            ['STATIONERY', '370866', '3', 'VOL'],
                            ['STATIONERY', '332641', '3', 'VOL'],
                            ['STATIONERY', '370871', '1', 'VOL'],
                            ['STATIONERY', '332633', '1', 'VOL'],
                            ['STATIONERY', '470201', '30', 'VOL'],
                            ['STATIONERY', '470216', '24', 'VOL'],
                            ['STATIONERY', '470237', '20', 'VOL'],
                            ['STATIONERY', '470245', '10', 'VOL'],
                            ['STATIONERY', '470282', '1', 'SET'],
                            ['STATIONERY', '470297', '1', 'BOX'],
                            ['STATIONERY', '470304', '10', 'PC'],
                            ['STATIONERY', '470311', '10', 'PC'],
                            ['STATIONERY', '470333', '10', 'PC'],
                            ['STATIONERY', '470334', '30', 'PC'],
                            ['STATIONERY', '470344', '5', 'PC'],
                            ['STATIONERY', '470364', '2', 'PKT'],
                            ['STATIONERY', '470375', '20', 'PC'],
                            ['STATIONERY', '470376', '20', 'PC'],
                            ['STATIONERY', '470377', '20', 'PC'],
                            ['STATIONERY', '470378', '20', 'PC'],
                            ['STATIONERY', '470379', '20', 'PC'],
                            ['STATIONERY', '470392', '10', 'VOL'],
                            ['STATIONERY', '470393', '5', 'VOL'],
                            ['STATIONERY', '470402', '10', 'PC'],
                            ['STATIONERY', '470405', '3', 'PC'],
                            ['STATIONERY', '470437', '1', 'PKT'],
                            ['STATIONERY', '470481', '10', 'PC'],
                            ['STATIONERY', '470524', '1', 'SET'],
                            ['STATIONERY', '470544', '2', 'PC'],
                            ['STATIONERY', '470547', '5', 'PC'],
                            ['STATIONERY', '470559', '1', 'BOT'],
                            ['STATIONERY', '470560', '1', 'BOT'],
                            ['STATIONERY', '470577', '7', 'PC'],
                            ['STATIONERY', '470578', '7', 'PC'],
                            ['STATIONERY', '470621', '10', 'PC'],
                            ['STATIONERY', '470622', '10', 'PC'],
                            ['STATIONERY', '470623', '10', 'PC'],
                            ['STATIONERY', '470641', '5', 'PC'],
                            ['STATIONERY', '470642', '5', 'PC'],
                            ['STATIONERY', '470643', '5', 'PC'],
                            ['STATIONERY', '470671', '10', 'PC'],
                            ['STATIONERY', '470672', '10', 'PC'],
                            ['STATIONERY', '470673', '10', 'PC'],
                            ['STATIONERY', '470674', '10', 'PC'],
                            ['STATIONERY', '470675', '10', 'PC'],
                            ['STATIONERY', '471002', '5', 'PC'],
                            ['STATIONERY', '471012', '10', 'PC'],
                            ['STATIONERY', '471013', '5', 'PC'],
                            ['STATIONERY', '471014', '5', 'PKT'],
                            ['STATIONERY', '471015', '2', 'PC'],
                            ['STATIONERY', '471022', '5', 'PC'],
                            ['STATIONERY', '471026', '1', 'PKT'],
                            ['STATIONERY', '471102', '5', 'PC'],
                            ['STATIONERY', '471106', '5', 'BOX'],
                            ['STATIONERY', '471108', '3', 'PC'],
                            ['STATIONERY', '471116', '10', 'BOX'],
                            ['STATIONERY', '471118', '10', 'BOX'],
                            ['STATIONERY', '471120', '1', 'BOX'],
                            ['STATIONERY', '471125', '20', 'PC'],
                            ['STATIONERY', '471126', '20', 'PC'],
                            ['STATIONERY', '471127', '20', 'PC'],
                            ['STATIONERY', '471129', '20', 'PC'],
                            ['STATIONERY', '471141', '30', 'PC'],
                            ['STATIONERY', '471143', '30', 'PC'],
                            ['STATIONERY', '471145', '30', 'PC'],
                            ['STATIONERY', '471151', '20', 'PC'],
                            ['STATIONERY', '471152', '20', 'PC'],
                            ['STATIONERY', '471155', '20', 'PC'],
                            ['STATIONERY', '471159', '1', 'PC'],
                            ['STATIONERY', '471167', '2', 'PC'],
                            ['STATIONERY', '471172', '2', 'PC'],
                            ['STATIONERY', '471262', '5', 'PC'],
                            ['STATIONERY', '471256', '3', 'PC'],
                            ['STATIONERY', '471273', '5', 'RL'],
                            ['STATIONERY', '471302', '2', 'PC'],
                            ['STATIONERY', '471321', '10', 'PKT'],
                            ['STATIONERY', '471322', '10', 'PC'],
                            ['STATIONERY', '471571', '1', 'PC'],
                            ['STATIONERY', '471645', '10', 'PC'],
                            ['STATIONERY', '471646', '10', 'PC'],
                            ['STATIONERY', '471648', '10', 'PC'],
                            ['STATIONERY', '471651', '20', 'PC'],
                            ['STATIONERY', '471652', '40', 'PC'],
                            ['STATIONERY', '471654', '20', 'PC'],
                            ['STATIONERY', '471671', '3', 'PC'],
                            ['STATIONERY', '471672', '3', 'PC'],
                            ['STATIONERY', '471673', '3', 'PC'],
                            ['STATIONERY', '471675', '3', 'PC'],
                            ['STATIONERY', '471809', '8', 'PC'],
                            ['STATIONERY', '473401', '1', 'SET'],
                            ['STATIONERY', '473412', '3', 'PC'],
                            ['STATIONERY', '473419', '3', 'PC'],
                            ['STATIONERY', '473415', '2', 'PC'],
                            ['STATIONERY', '473417', '2', 'PC'],
                            ['STATIONERY', '473418', '2', 'PC'],
                            ['STATIONERY', '473422', '2', 'PC'],
                            ['STATIONERY', '473423', '2', 'PC'],
                            ['STATIONERY', '473424', '2', 'PC'],
                            ['STATIONERY', '473425', '2', 'PC']
            ]
    qtydup_table_stw = [
                            ['dept', 'impa', 'qty', 'unit'],
                            ['STEWARD', '150616', '24', 'SHT'],
                            ['STEWARD', '171090', '20', 'PKT'],
                            ['STEWARD', '171454', '1', 'CTN'],
                            ['STEWARD', '173404', '5', 'PKT'],
                            ['STEWARD', '174058', '24', 'PC'],
                            ['STEWARD', '174072', '24', 'PC'],
                            ['STEWARD', '174175', '10', 'PKT'],
                            ['STEWARD', '174175', '50', 'PKT'],
                            ['STEWARD', '174202', '10', 'PC'],
                            ['STEWARD', '174206', '30', 'PC'],
                            ['STEWARD', '174220', '150', 'BOX'],
                            ['STEWARD', '174241', '10', 'CTN'],
                            ['STEWARD', '174275', '2', 'PC'],
                            ['STEWARD', '174277', '6', 'PC'],
                            ['STEWARD', '190122', '4', 'PR'],
                            ['STEWARD', '550108', '10', 'CTN'],
                            ['STEWARD', '550111', '100', 'PKT'],
                            ['STEWARD', '550121', '18', 'BTL'],
                            ['STEWARD', '550124', '30', 'BTL'],
                            ['STEWARD', '550133', '3', 'BAG'],
                            ['STEWARD', '550141', '50', 'BTL'],
                            ['STEWARD', '550163', '10', 'BTL'],
                            ['STEWARD', '550165', '20', 'TIN'],
                            ['STEWARD', '550171', '24', 'BTL'],
                            ['STEWARD', '550254', '250', 'PC'],
                            ['STEWARD', '550301', '10', 'TIN'],
                            ['STEWARD', '550307', '20', 'TIN'],
                            ['STEWARD', '550331', '10', 'TIN'],
                            ['STEWARD', '550336', '50', 'PC'],
                            ['STEWARD', '110903', '1', 'SET'],
                            ['STEWARD', '150101', '25', 'PC'],
                            ['STEWARD', '150142', '2', 'PC'],
                            ['STEWARD', '150281', '25', 'PC'],
                            ['STEWARD', '150286', '25', 'PC'],
                            ['STEWARD', '150461', '2', 'PC'],
                            ['STEWARD', '150464', '2', 'PC'],
                            ['STEWARD', '150527', '100', 'PC'],
                            ['STEWARD', '150602', '24', 'SHT'],
                            ['STEWARD', '150607', '24', 'SHT'],
                            ['STEWARD', '150617', '10', 'SHT'],
                            ['STEWARD', '150652', '1', 'RL'],
                            ['STEWARD', '170101', '25', 'PC'],
                            ['STEWARD', '170102', '25', 'PC'],
                            ['STEWARD', '170103', '25', 'PC'],
                            ['STEWARD', '170104', '25', 'PC'],
                            ['STEWARD', '170118', '25', 'PC'],
                            ['STEWARD', '170169', '25', 'PC'],
                            ['STEWARD', '170170', '25', 'PC'],
                            ['STEWARD', '170173', '25', 'PC'],
                            ['STEWARD', '170311', '25', 'PC'],
                            ['STEWARD', '170315', '25', 'PC'],
                            ['STEWARD', '170317', '25', 'PC'],
                            ['STEWARD', '170329', '25', 'PC'],
                            ['STEWARD', '170331', '20', 'PC'],
                            ['STEWARD', '170363', '10', 'PC'],
                            ['STEWARD', '170402', '25', 'PC'],
                            ['STEWARD', '170403', '25', 'PC'],
                            ['STEWARD', '170404', '25', 'PC'],
                            ['STEWARD', '170405', '25', 'PC'],
                            ['STEWARD', '170409', '25', 'PC'],
                            ['STEWARD', '170413', '25', 'PC'],
                            ['STEWARD', '170433', '2', 'PC'],
                            ['STEWARD', '170605', '25', 'PC'],
                            ['STEWARD', '170712', '2', 'PC'],
                            ['STEWARD', '170756', '4', 'PC'],
                            ['STEWARD', '170793', '4', 'PC'],
                            ['STEWARD', '170794', '8', 'PC'],
                            ['STEWARD', '170805', '4', 'PC'],
                            ['STEWARD', '170822', '4', 'PC'],
                            ['STEWARD', '171026', '4', 'SET'],
                            ['STEWARD', '171033', '2', 'SET'],
                            ['STEWARD', '171047', '6', 'PC'],
                            ['STEWARD', '171051', '6', 'PC'],
                            ['STEWARD', '171106', '4', 'PC'],
                            ['STEWARD', '171122', '6', 'PC'],
                            ['STEWARD', '171218', '4', 'PC'],
                            ['STEWARD', '171351', '2', 'SET'],
                            ['STEWARD', '171357', '2', 'PC'],
                            ['STEWARD', '171372', '2', 'PC'],
                            ['STEWARD', '171425', '2', 'PC'],
                            ['STEWARD', '171463', '4', 'PC'],
                            ['STEWARD', '171708', '2', 'PC'],
                            ['STEWARD', '171713', '1', 'PC'],
                            ['STEWARD', '171805', '1', 'PC'],
                            ['STEWARD', '171806', '1', 'PC'],
                            ['STEWARD', '171824', '1', 'PC'],
                            ['STEWARD', '171825', '1', 'PC'],
                            ['STEWARD', '171828', '1', 'PC'],
                            ['STEWARD', '171842', '1', 'PC'],
                            ['STEWARD', '171844', '1', 'PC'],
                            ['STEWARD', '171849', '1', 'PC'],
                            ['STEWARD', '171879', '1', 'PC'],
                            ['STEWARD', '171922', '1', 'PC'],
                            ['STEWARD', '172011', '1', 'PC'],
                            ['STEWARD', '172013', '1', 'PC'],
                            ['STEWARD', '172023', '2', 'PC'],
                            ['STEWARD', '172026', '2', 'PC'],
                            ['STEWARD', '172121', '2', 'PC'],
                            ['STEWARD', '172125', '2', 'PC'],
                            ['STEWARD', '172158', '1', 'PC'],
                            ['STEWARD', '172162', '2', 'PC'],
                            ['STEWARD', '172166', '2', 'PC'],
                            ['STEWARD', '172207', '1', 'PC'],
                            ['STEWARD', '172219', '2', 'PC'],
                            ['STEWARD', '172226', '2', 'PC'],
                            ['STEWARD', '172253', '2', 'PC'],
                            ['STEWARD', '172262', '1', 'PC'],
                            ['STEWARD', '172265', '1', 'PC'],
                            ['STEWARD', '172302', '1', 'PC'],
                            ['STEWARD', '172312', '1', 'PC'],
                            ['STEWARD', '172317', '1', 'PC'],
                            ['STEWARD', '172318', '1', 'PC'],
                            ['STEWARD', '172323', '1', 'PC'],
                            ['STEWARD', '172342', '2', 'PC'],
                            ['STEWARD', '172347', '2', 'PC'],
                            ['STEWARD', '172362', '2', 'PC'],
                            ['STEWARD', '172370', '2', 'PC'],
                            ['STEWARD', '172372', '2', 'PC'],
                            ['STEWARD', '172383', '1', 'PC'],
                            ['STEWARD', '172386', '3', 'PC'],
                            ['STEWARD', '172401', '1', 'PC'],
                            ['STEWARD', '172404', '1', 'PC'],
                            ['STEWARD', '172406', '1', 'PC'],
                            ['STEWARD', '172433', '2', 'PC'],
                            ['STEWARD', '172436', '2', 'PC'],
                            ['STEWARD', '172504', '2', 'PC'],
                            ['STEWARD', '172512', '25', 'PC'],
                            ['STEWARD', '172541', '1', 'PC'],
                            ['STEWARD', '172548', '2', 'PC'],
                            ['STEWARD', '172555', '2', 'PC'],
                            ['STEWARD', '172561', '2', 'PC'],
                            ['STEWARD', '172562', '2', 'PC'],
                            ['STEWARD', '172563', '2', 'PC'],
                            ['STEWARD', '172573', '2', 'PC'],
                            ['STEWARD', '172600', '6', 'PC'],
                            ['STEWARD', '172611', '2', 'PC'],
                            ['STEWARD', '172617', '4', 'PC'],
                            ['STEWARD', '172716', '6', 'PC'],
                            ['STEWARD', '172728', '1', 'PC'],
                            ['STEWARD', '172741', '1', 'PC'],
                            ['STEWARD', '172774', '1', 'PC'],
                            ['STEWARD', '172802', '1', 'PC'],
                            ['STEWARD', '172812', '1', 'PC'],
                            ['STEWARD', '172830', '1', 'PC'],
                            ['STEWARD', '172856', '2', 'PC'],
                            ['STEWARD', '172872', '1', 'PC'],
                            ['STEWARD', '172873', '1', 'PC'],
                            ['STEWARD', '172936', '3', 'PC'],
                            ['STEWARD', '172937', '3', 'PC'],
                            ['STEWARD', '172938', '3', 'PC'],
                            ['STEWARD', '172971', '6', 'PC'],
                            ['STEWARD', '173128', '25', 'PC'],
                            ['STEWARD', '173129', '25', 'PC'],
                            ['STEWARD', '173131', '25', 'PC'],
                            ['STEWARD', '173132', '25', 'PC'],
                            ['STEWARD', '173133', '25', 'PC'],
                            ['STEWARD', '173141', '6', 'PC'],
                            ['STEWARD', '173142', '4', 'PC'],
                            ['STEWARD', '173172', '1', 'PC'],
                            ['STEWARD', '173176', '1', 'PC'],
                            ['STEWARD', '173202', '1', 'PC'],
                            ['STEWARD', '173216', '2', 'PC'],
                            ['STEWARD', '173226', '2', 'PC'],
                            ['STEWARD', '173238', '1', 'PC'],
                            ['STEWARD', '173258', '2', 'PC'],
                            ['STEWARD', '173260', '2', 'PC'],
                            ['STEWARD', '173268', '2', 'PC'],
                            ['STEWARD', '173273', '2', 'PC'],
                            ['STEWARD', '173277', '2', 'PC'],
                            ['STEWARD', '173301', '2', 'PC'],
                            ['STEWARD', '173336', '2', 'PC'],
                            ['STEWARD', '173341', '2', 'PC'],
                            ['STEWARD', '173363', '1', 'PC'],
                            ['STEWARD', '173366', '1', 'PC'],
                            ['STEWARD', '173371', '1', 'PC'],
                            ['STEWARD', '173381', '1', 'PC'],
                            ['STEWARD', '173412', '3', 'PR'],
                            ['STEWARD', '173471', '2', 'PC'],
                            ['STEWARD', '173636', '1', 'PC'],
                            ['STEWARD', '173638', '1', 'PC'],
                            ['STEWARD', '173651', '1', 'PC'],
                            ['STEWARD', '174005', '1', 'PC'],
                            ['STEWARD', '174011', '1', 'PC'],
                            ['STEWARD', '174029', '2', 'PC'],
                            ['STEWARD', '174048', '2', 'PR'],
                            ['STEWARD', '174051', '2', 'PC'],
                            ['STEWARD', '174052', '2', 'PC'],
                            ['STEWARD', '174055', '20', 'PC'],
                            ['STEWARD', '174124', '5', 'PC'],
                            ['STEWARD', '174142', '25', 'PC'],
                            ['STEWARD', '174150', '25', 'PC'],
                            ['STEWARD', '174163', '3', 'PC'],
                            ['STEWARD', '174252', '20', 'PC'],
                            ['STEWARD', '174256', '4', 'PC'],
                            ['STEWARD', '174276', '6', 'PC'],
                            ['STEWARD', '174280', '2', 'PC'],
                            ['STEWARD', '174303', '2', 'BOX'],
                            ['STEWARD', '174316', '10', 'PC'],
                            ['STEWARD', '174317', '10', 'PC'],
                            ['STEWARD', '174356', '4', 'PC'],
                            ['STEWARD', '174521', '2', 'PC'],
                            ['STEWARD', '174556', '2', 'PC'],
                            ['STEWARD', '174562', '2', 'PC'],
                            ['STEWARD', '174629', '2', 'SET'],
                            ['STEWARD', '174671', '2', 'SET'],
                            ['STEWARD', '174721', '1', 'SET'],
                            ['STEWARD', '174728', '1', 'PC'],
                            ['STEWARD', '175091', '2', 'SET'],
                            ['STEWARD', '330950', '1', 'SET'],
                            ['STEWARD', '350206', '20', 'MTR'],
                            ['STEWARD', '470245', '1', 'PC'],
                            ['STEWARD', '471517', '4', 'PC'],
                            ['STEWARD', '510611', '2', 'PC'],
                            ['STEWARD', '510646', '1', 'PC'],
                            ['STEWARD', '510801', '5', 'PC'],
                            ['STEWARD', '510806', '10', 'PC'],
                            ['STEWARD', '511008', '6', 'SHT'],
                            ['STEWARD', '511023', '5', 'SHT'],
                            ['STEWARD', '511036', '10', 'SHT'],
                            ['STEWARD', '550322', '5', 'PC'],
                            ['STEWARD', '614010', '2', 'PC']
            ]
    qtydup_table_med = [
                            ['dept', 'impa', 'qty', 'unit'],
                            ['MEDICAL', '391701', '300', 'PC'],
                            ['MEDICAL', '391702', '1', 'BOX'],
                            ['MEDICAL', '391703', '1', 'BOX'],
                            ['MEDICAL', '391704', '6', 'BOX'],
                            ['MEDICAL', '391705', '30', 'PC'],
                            ['MEDICAL', '391706', '5', 'PC'],
                            ['MEDICAL', '391707', '2', 'BOX'],
                            ['MEDICAL', '391708', '2', 'BOX'],
                            ['MEDICAL', '391709', '2', 'BOX'],
                            ['MEDICAL', '391710', '45', 'BTL'],
                            ['MEDICAL', '391711', '1', 'BOX'],
                            ['MEDICAL', '391712', '2', 'BTL'],
                            ['MEDICAL', '391713', '4', 'BOX'],
                            ['MEDICAL', '391714', '1', 'BTL'],
                            ['MEDICAL', '391715', '5', 'BOX'],
                            ['MEDICAL', '390462', '2', 'BTL'],
                            ['MEDICAL', '391717', '40', 'PC'],
                            ['MEDICAL', '390631', '4', 'BOX'],
                            ['MEDICAL', '391719', '6', 'BTL'],
                            ['MEDICAL', '391720', '2', 'BTL'],
                            ['MEDICAL', '391721', '1', 'BOX'],
                            ['MEDICAL', '391722', '1', 'BOX'],
                            ['MEDICAL', '391723', '1', 'BTL'],
                            ['MEDICAL', '391724', '2', 'BOX'],
                            ['MEDICAL', '390644', '6', 'TUB'],
                            ['MEDICAL', '391726', '150', 'PC'],
                            ['MEDICAL', '391727', '4', 'BOX'],
                            ['MEDICAL', '391728', '2', 'BOX'],
                            ['MEDICAL', '391729', '3', 'BOX'],
                            ['MEDICAL', '391730', '20', 'PC'],
                            ['MEDICAL', '391731', '2', 'BOX'],
                            ['MEDICAL', '391732', '1', 'BOX'],
                            ['MEDICAL', '390652', '6', 'TUB'],
                            ['MEDICAL', '391734', '2', 'BOX'],
                            ['MEDICAL', '391735', '1', 'BOX'],
                            ['MEDICAL', '390390', '6', 'BOX'],
                            ['MEDICAL', '391737', '1', 'BOX'],
                            ['MEDICAL', '390656', '2', 'BOX'],
                            ['MEDICAL', '391739', '60', 'PC'],
                            ['MEDICAL', '391740', '2', 'BOX'],
                            ['MEDICAL', '390658', '20', 'BTL'],
                            ['MEDICAL', '391742', '6', 'BTL'],
                            ['MEDICAL', '390660', '3', 'BOX'],
                            ['MEDICAL', '391744', '2', 'BTL'],
                            ['MEDICAL', '391745', '2', 'BTL'],
                            ['MEDICAL', '391746', '4', 'TUB'],
                            ['MEDICAL', '391748', '4', 'BTL'],
                            ['MEDICAL', '391747', '60', 'PC'],
                            ['MEDICAL', '390472', '4', 'BTL'],
                            ['MEDICAL', '390672', '3', 'BTL'],
                            ['MEDICAL', '391718', '1', 'PC'],
                            ['MEDICAL', '390673', '5', 'BTL'],
                            ['MEDICAL', '390674', '2', 'BTL'],
                            ['MEDICAL', '391750', '6', 'TUB'],
                            ['MEDICAL', '390678', '1', 'BOX'],
                            ['MEDICAL', '390679', '1', 'PC'],
                            ['MEDICAL', '391725', '50', 'PC'],
                            ['MEDICAL', '391751', '1', 'BOX'],
                            ['MEDICAL', '390682', '8', 'BOX'],
                            ['MEDICAL', '391753', '1', 'BTL'],
                            ['MEDICAL', '390683', '5', 'TUB'],
                            ['MEDICAL', '391801', '1', 'SET'],
                            ['MEDICAL', '391802', '1', 'SET'],
                            ['MEDICAL', '391803', '1', 'SET'],
                            ['MEDICAL', '391804', '1', 'PC'],
                            ['MEDICAL', '391805', '1', 'PC'],
                            ['MEDICAL', '391806', '1', 'PC'],
                            ['MEDICAL', '391834', '1', 'PC'],
                            ['MEDICAL', '391807', '1', 'SET'],
                            ['MEDICAL', '391835', '2', 'SET'],
                            ['MEDICAL', '391808', '1', 'SET'],
                            ['MEDICAL', '391809', '1', 'SET'],
                            ['MEDICAL', '391810', '2', 'BOX'],
                            ['MEDICAL', '391811', '10', 'PKT'],
                            ['MEDICAL', '391812', '10', 'PKT'],
                            ['MEDICAL', '391813', '1', 'ROL'],
                            ['MEDICAL', '391814', '3', 'PKT'],
                            ['MEDICAL', '391815', '4', 'BOX'],
                            ['MEDICAL', '391816', '3', 'ROL'],
                            ['MEDICAL', '391817', '10', 'PKT'],
                            ['MEDICAL', '391836', '10', 'PKT'],
                            ['MEDICAL', '391843', '10', 'PKT'],
                            ['MEDICAL', '391818', '1', 'BOX'],
                            ['MEDICAL', '391844', '1', 'PC'],
                            ['MEDICAL', '391819', '10', 'ROL'],
                            ['MEDICAL', '391820', '5', 'PC'],
                            ['MEDICAL', '391821', '4', 'PC'],
                            ['MEDICAL', '391822', '2', 'BTL'],
                            ['MEDICAL', '391823', '6', 'ROL'],
                            ['MEDICAL', '391824', '1', 'PKT'],
                            ['MEDICAL', '391351', '2', 'BOX'],
                            ['MEDICAL', '391826', '50', 'PC'],
                            ['MEDICAL', '391827', '1', 'PC'],
                            ['MEDICAL', '391828', '6', 'PC'],
                            ['MEDICAL', '391845', '6', 'PC'],
                            ['MEDICAL', '391829', '6', 'PC'],
                            ['MEDICAL', '391830', '1', 'BOX'],
                            ['MEDICAL', '391874', '1', 'BOX'],
                            ['MEDICAL', '391831', '5', 'PRS'],
                            ['MEDICAL', '391863', '5', 'PRS'],
                            ['MEDICAL', '391866', '5', 'PRS'],
                            ['MEDICAL', '391250', '3', 'PC'],
                            ['MEDICAL', '391832', '10', 'PC'],
                            ['MEDICAL', '391833', '1', 'PC'],
                            ['MEDICAL', '391361', '1', 'PC'],
                            ['MEDICAL', '391362', '1', 'PC'],
                            ['MEDICAL', '391269', '1', 'PC'],
                            ['MEDICAL', '391837', '1', 'PC'],
                            ['MEDICAL', '391838', '1', 'PC'],
                            ['MEDICAL', '391839', '2', 'PC'],
                            ['MEDICAL', '391840', '10', 'PKT'],
                            ['MEDICAL', '391841', '100', 'PC'],
                            ['MEDICAL', '391842', '2', 'BTL'],
                            ['MEDICAL', '391334', '1', 'PKT'],
                            ['MEDICAL', '391371', '1', 'PC'],
                            ['MEDICAL', '391368', '1', 'SET'],
                            ['MEDICAL', '391846', '3', 'PC'],
                            ['MEDICAL', '391847', '1', 'PC'],
                            ['MEDICAL', '391848', '1', 'PC'],
                            ['MEDICAL', '391849', '1', 'SET'],
                            ['MEDICAL', '391869', '1', 'PC'],
                            ['MEDICAL', '391253', '1', 'PC'],
                            ['MEDICAL', '391420', '1', 'PC'],
                            ['MEDICAL', '391870', '10', 'PC'],
                            ['MEDICAL', '391871', '1', 'PC'],
                            ['MEDICAL', '391872', '1', 'PC'],
                            ['MEDICAL', '370674', '1', 'PC'],
                            ['MEDICAL', '391888', '1', 'PC'],
                            ['MEDICAL', '391889', '1', 'PC'],
                            ['MEDICAL', '391890', '1', 'PC'],
                            ['MEDICAL', '391891', '1', 'PC'],
                            ['MEDICAL', '370682', '1', 'PC'],
                            ['MEDICAL', '391850', '30', 'PC'],
                            ['MEDICAL', '391851', '30', 'PC'],
                            ['MEDICAL', '391893', '20', 'PC'],
                            ['MEDICAL', '391852', '40', 'PC'],
                            ['MEDICAL', '391853', '40', 'PC'],
                            ['MEDICAL', '391854', '40', 'PC'],
                            ['MEDICAL', '391855', '10', 'PC'],
                            ['MEDICAL', '391894', '10', 'PC'],
                            ['MEDICAL', '391856', '10', 'PC'],
                            ['MEDICAL', '391857', '1', 'PC'],
                            ['MEDICAL', '391858', '1', 'PC'],
                            ['MEDICAL', '391895', '1', 'PC'],
                            ['MEDICAL', '391859', '2', 'PC'],
                            ['MEDICAL', '391896', '2', 'PC'],
                            ['MEDICAL', '391897', '2', 'PC'],
                            ['MEDICAL', '391860', '2', 'PC'],
                            ['MEDICAL', '391861', '2', 'PC'],
                            ['MEDICAL', '391862', '10', 'PC'],
                            ['MEDICAL', '391316', '10', 'PC'],
                            ['MEDICAL', '391864', '1', 'PKT'],
                            ['MEDICAL', '391865', '1', 'PC'],
                            ['MEDICAL', '391283', '2', 'BOX'],
                            ['MEDICAL', '191867', '1', 'PC'],
                            ['MEDICAL', '391868', '1', 'PC'],
                            ['MEDICAL', '391208', '1', 'PC'],
                            ['MEDICAL', '391290', '1', 'PC'],
                            ['MEDICAL', '391440', '1', 'PC'],
                            ['MEDICAL', '391292', '1', 'PC'],
                            ['MEDICAL', '391873', '1', 'PC'],
                            ['MEDICAL', '391226', '300', 'PC'],
                            ['MEDICAL', '391875', '1', 'PC'],
                            ['MEDICAL', '391876', '1', 'PC'],
                            ['MEDICAL', '391877', '1', 'PC'],
                            ['MEDICAL', '391879', '1', 'PC'],
                            ['MEDICAL', '391880', '30', 'PC'],
                            ['MEDICAL', '391881', '6', 'PC'],
                            ['MEDICAL', '391898', '6', 'PC'],
                            ['MEDICAL', '391882', '1', 'ROL'],
                            ['MEDICAL', '391899', '1', 'ROL'],
                            ['MEDICAL', '391883', '2', 'ROL'],
                            ['MEDICAL', '391884', '100', 'PC'],
                            ['MEDICAL', '391885', '1', 'PC'],
                            ['MEDICAL', '651715', '1', 'PC'],
                            ['MEDICAL', '391929', '1', 'PC'],
                            ['MEDICAL', '391930', '1', 'PC'],
                            ['MEDICAL', '391952', '1', 'PC'],
                            ['MEDICAL', '391945', '1', 'PC'],
                            ['MEDICAL', '391946', '2', 'PC'],
                            ['MEDICAL', '391947', '20', 'PC'],
                            ['MEDICAL', '391948', '10', 'PC'],
                            ['MEDICAL', '391949', '3', 'PC'],
                            ['MEDICAL', '391886', '1', 'PC'],
                            ['MEDICAL', '391887', '2', 'PC'],
                            ['MEDICAL', '391950', '2', 'PC'],
                            ['MEDICAL', '391391', '1', 'PC'],
                            ['MEDICAL', '391951', '1', 'PC']
            ]
    df_qtytable_deck = pd.DataFrame(qtydup_table_deck[1:], columns=qtydup_table_deck[0])
    df_qtytable_eng = pd.DataFrame(qtydup_table_eng[1:], columns=qtydup_table_eng[0])
    df_qtytable_sta = pd.DataFrame(qtydup_table_sta[1:], columns=qtydup_table_sta[0])
    df_qtytable_stw = pd.DataFrame(qtydup_table_stw[1:], columns=qtydup_table_stw[0])
    df_qtytable_med = pd.DataFrame(qtydup_table_med[1:], columns=qtydup_table_med[0])
    
    # ---- process for auro remark & duplicated check ----
    for row in range(2,row_count-1):
        article_data = combined_data[row][order.index("article_col")]
        if article_data is not None:
            # duplicated check
            base_dept = combined_data[row][order.index("dept_col")]
            base_no = combined_data[row][order.index("no_col")]
            base_article = combined_data[row][order.index("article_col")]
            for row_dup in range(row + 1, row_count):
                dup_article = combined_data[row_dup][order.index("article_col")] 
                if dup_article is not None:
                    dup_dept = combined_data[row_dup][order.index("dept_col")]
                    dup_article = combined_data[row_dup][order.index("article_col")]
                    if base_dept == dup_dept and base_article == dup_article and base_no is not None:
                        remark_content[row_dup] = "Duplicated w/" + str(base_dept) + "-" + str(base_no)

            # search the remark target
            remark_flag = False
            for index, df_row in df_remarkbase_1.iterrows():
                content = df_row['CONTENT']
                if not isinstance(content, int) and content in article_data:
                    if not remark_content[row]:
                        remark_content[row] = str(df_row['REMARK'])
                    elif remark_content[row]:
                        remark_content[row] = str(remark_content[row]) + " / " + str(df_row['REMARK'])
                    remark_flag = True
                    break

            if not remark_flag:
                for index_2, df_row_2 in df_remarkbase_2.iterrows():
                    content_2 = df_row_2['CONTENT']
                    if not isinstance(content_2, int) and content_2 in article_data:
                        if not remark_content[row]:
                            remark_content[row] = str(df_row_2['REMARK'])
                        elif remark_content[row]:
                            remark_content[row] = str(remark_content[row]) + " / " + str(df_row_2['REMARK'])
                        remark_flag = True
                        break

    remark_content.insert(0, "") # length adjust with excel
#   st.write(remark_content)####


    # ---- search the too much QTY ----
    def check_quantity(base_dept, df_qtytable):
        for index, df_row in df_qtytable.iterrows():
            if ' ' in df_row['qty']: # skip the half-space
                continue
            else:
                if df_row['impa'] in base_impa and df_row['unit'] in base_unit and (int(df_row['qty']) * 3) < int(base_qty):
                    if combined_data[row][order.index("remark_col")] is None:
                        combined_data[row][order.index("remark_col")] = "too much qty"
                    else:
                        combined_data[row][order.index("remark_col")] += "too much qty"
                    break


    # ---- QTY check process ----
    for row in range(2,row_count):
        base_dept = combined_data[row][order.index("dept_col")]
        base_impa = combined_data[row][order.index("impa_col")]
        base_qty = combined_data[row][order.index("qty_col")]
        base_unit = combined_data[row][order.index("unit_col")]
        if base_dept is not None and base_impa is not None and base_qty is not None and base_unit is not None:
            if base_dept == "DECK":
                check_quantity(base_dept, df_qtytable_deck)
            elif base_dept == "ENGINE":
                check_quantity(base_dept, df_qtytable_eng)
            elif base_dept == "STATIONERY":
                check_quantity(base_dept, df_qtytable_sta)
            elif base_dept == "STEWARD":
                check_quantity(base_dept, df_qtytable_stw)
            elif base_dept == "MEDICAL":
                check_quantity(base_dept, df_qtytable_med)

    return combined_data, remark_content


# Excel content control
def excelcontrol(combined_data, sheet, auto_remark_flag, remark_content, mode):
    # mode call
    content_data, order, page_col, column_widths, font, font_size, zoom = mode_select(mode)

    # ---- advance setting ----
    for col_idx, cell_value in enumerate(content_data, start=1):
        # Header input
        col_letter = get_column_letter(col_idx)
        sheet[f"{col_letter}1"].value = cell_value
        sheet[f"{col_letter}1"].font = Font(name=font, size=font_size)
        # Column width setting
        sheet.column_dimensions[col_letter].width = column_widths[col_idx - 1]


    # font & size control
    for column in sheet.iter_cols(min_col=0, max_col=len(combined_data[0])):
        for cell in column:
            cell.font = Font(name=font, size=font_size)


    # ---- 1st for next ----
    # price digit analysis
    digit_flag = False
    for row in sheet.iter_rows(min_row=2, max_row=sheet.max_row):
        if "." in str(sheet.cell(row=row[0].row, column=content_data.index("S_U_PRICE")+1).value):
            digit_flag = True
            break


    # ---- 2nd for next ----
    # color set
    fill_ye = PatternFill(start_color="FFFF00", end_color="FFFF00", fill_type="solid") # Yellow
    fill_or = PatternFill(patternType='solid', fgColor='FFC000') # Orange
    fill_gr = PatternFill(fill_type="solid", fgColor="C6EFCE") #Green
    for row in sheet.iter_rows(min_row=2, max_row=sheet.max_row, min_col=1, max_col=sheet.max_column):
        # variable set
        dept_cell = sheet.cell(row=row[0].row, column=content_data.index("DEPT")+1)
        no_cell = sheet.cell(row=row[0].row, column=content_data.index("NOS")+1)
        impa_cell = sheet.cell(row=row[0].row, column=content_data.index("IMPA")+1)
        article_cell = sheet.cell(row=row[0].row, column=content_data.index("ARTICLE")+1)
        qty_cell = sheet.cell(row=row[0].row, column=content_data.index("QTY")+1)
        unit_cell = sheet.cell(row=row[0].row, column=content_data.index("UNIT")+1)
        s_u_price_cell = sheet.cell(row=row[0].row, column=content_data.index("S_U_PRICE")+1)
        s_amount_cell = sheet.cell(row=row[0].row, column=content_data.index("S_AMOUNT")+1)
        p_u_price_cell = sheet.cell(row=row[0].row, column=content_data.index("P_U_PRICE")+1)
        p_amount_cell = sheet.cell(row=row[0].row, column=content_data.index("P_AMOUNT")+1)
        remark_cell = sheet.cell(row=row[0].row, column=content_data.index("REMARK")+1)
        auto_remark_cell = sheet.cell(row=row[0].row, column=order.index("auto_remark_col")+1)

        # sales amount formula
        if s_amount_cell.value is not None:
            if qty_cell.value is not None and s_u_price_cell.value is not None:
                s_amount_cell.value = f"=F{row[0].row}*H{row[0].row}"
            else:
                s_amount_cell.value = ""
            # remark : previous process is "isinstance.value,(int, float)" but had a problem with accuracy.
#            try:
#                value = float(s_u_price_cell.value)
#                s_amount_cell.value = f"=F{row[0].row}*H{row[0].row}"
#            except (ValueError, TypeError):
#                s_amount_cell.value = None
        # Align price to last two digits
        if digit_flag:
            s_u_price_cell.number_format = s_amount_cell.number_format = '0.00'

        # purchase amount formula
        if "P_AMOUNT" in content_data and p_amount_cell.value is not None:
            if qty_cell.value is not None and p_u_price_cell.value is not None:
                p_amount_cell.value = f'=F{row[0].row}*L{row[0].row}'
            else:
                p_amount_cell.value = ""
            # Align price to last two digits
            if digit_flag:
                p_u_price_cell.number_format = p_amount_cell.number_format = '0.00'

        # date format control as */*
        if mode == "jp_store" or mode == "jp_spare":
            if isinstance(row[10].value, (int, float)):
                month = int(datetime.strptime(str(row[10].value), "%m").strftime("%-m"))
                day = int(datetime.strptime(str(row[10].value), "%d").strftime("%-d"))
                date_value = datetime(1900, month, day).strftime("%-m/%-d")
            else:
                date_value = None
                sheet[get_column_letter(row[10].column) + str(row[10].row)].number_format = 'm/d'
                sheet[get_column_letter(row[10].column) + str(row[10].row)].value = date_value
               
        # Color support
        # Yellow
        if qty_cell.value == 0 or qty_cell.value == str(0):
            qty_cell.fill = fill_ye
            if s_u_price_cell.value is not None:
                s_amount_cell.fill = fill_or
        elif article_cell.value is not None and qty_cell.value is None and unit_cell.value is not None:
            qty_cell.fill = fill_ye
            s_amount_cell.fill = fill_ye
        elif s_u_price_cell.value == 0:
            s_u_price_cell.fill = fill_ye
            s_amount_cell.fill = fill_ye
        # Orange
        if article_cell.value is not None and qty_cell.value is not None \
           and s_u_price_cell.value is not None and s_amount_cell.value is None:
            s_amount_cell.fill = fill_or
        elif s_amount_cell.value == 0 or s_amount_cell.value == "":
            s_amount_cell.fill = fill_or
        # Green
        if auto_remark_flag:
            if remark_content[row[0].row] is not None and remark_content[row[0].row] != []:
                auto_remark_cell.value = str(remark_content[row[0].row])
                auto_remark_cell.fill = fill_gr
                        
        # Auto remark support - too much qty as green
        if auto_remark_flag:
            remark_cell = remark_cell
            if remark_cell.value is not None and "too much qty" in str(remark_cell.value):
                qty_cell.fill = fill_gr
                remark_cell.value = remark_cell.value.replace("too much qty", "")

        # Alignment setting
        # Dept
        dept_cell.alignment = Alignment(vertical='top', horizontal='left')
        # No
        no_value = no_cell.value
        if isinstance(no_value, (int, float)) or (isinstance(no_value, str) and no_value.replace(".", "").isdigit()):
            no_cell.alignment = Alignment(vertical='top', horizontal='right')
        else:
            no_cell.alignment = Alignment(vertical='top', horizontal='left')
        # IMPA
        if mode == "jp_spare":
            sheet.cell(row=row[0].row, column=content_data.index("UNIT / MODEL")+1).alignment = Alignment(vertical='top', horizontal='left', wrap_text=True)
        else:
            impa_cell.alignment = Alignment(vertical='top', horizontal='left', wrap_text=True)
        # Article
        article_cell.alignment = Alignment(vertical='top', wrap_text=True)
        # QTY
        qty_cell.alignment = Alignment(vertical='top', horizontal='right')
        # Unit
        unit_cell.alignment = Alignment(vertical='top', horizontal='left')
        # S_U_PRICE
        s_u_price_cell.alignment = Alignment(vertical='top', horizontal='right')
        # S_AMOUNT
        s_amount_cell.alignment = Alignment(vertical='top', horizontal='right')
        # Remark
        remark_cell.alignment = Alignment(vertical='top', wrap_text=True)
            
#        sht.name & page_num clear
#        if sheet.cell(row=row[0].row, column=order.index("page_col")+1).value is not None:
#            sheet.cell(row=row[0].row, column=order.index("page_col")+1).value = None

    return sheet


# Main as Streamlit process
def main():
    # ---- Pass-code input process ----
    # load secrets case select
    password =  "093"#### for run with cmd
#    password = st.secrets["password"]#### for streamlit cloud

    # pass input
    login_tab = st.empty()
    login_tab.tabs(["Login"])
    user, user_placeholder, code, code_placeholder = security_code_input()

    # ---- Login process ----
    if code == password and user == "MASHIN":
        # login text_input erase & tab set
        login_tab.empty()
        user_placeholder.empty()
        code_placeholder.empty()
        tab1, tab2, tab3, tab4 = st.tabs(["Converter", "PDF", "manual", "tips"])

        # ---- converter process ----
        with tab1:
            # ---- Sidebar contorl ----
            # version info
            st.sidebar.caption("Version : 0.6.3")
            prog_info = st.sidebar.title("Please upload your file")
            st.sidebar.markdown("---")

            # mode select
            content_data = order = column_widths = []
            page_col = 0
            pagelist = ["Mashin system Store mode","Mashin system Spare mode","Netherlands sample","BUSAN sample"]
            selector=st.sidebar.selectbox("Mode/Country select",pagelist)
            if selector=="Mashin system Store mode":
                mode = "jp_store"
            elif selector=="Mashin system Spare mode":
                mode = "jp_spare"
            elif selector=="Netherlands sample":
                mode = "Netherlands"
            elif selector=="BUSAN sample":
                mode = "BUSAN"
            content_data, order, page_col, column_widths,font,font_size,zoom = mode_select(mode)
            st.sidebar.markdown("---")

            # ---- All sheet process select ----
            pagelist = ["Process only target sheet","Process to All sheets in the file"]
            selector=st.sidebar.selectbox("Sheet process select",pagelist)
            if selector=="Process only target sheet":
                all_sht_flag = False
            elif selector=="Process to All sheets in the file":
                all_sht_flag = True
            st.sidebar.markdown("---")

            # ---- Auto remark process select ----
            pagelist = ["No use Auto-remark","Apply Auto-remark"]
            selector=st.sidebar.selectbox("Auto remark process select",pagelist)
            if selector=="No use Auto-remark":
                auto_remark_flag = False
            elif selector=="Apply Auto-remark":
                auto_remark_flag = True
            st.sidebar.markdown("---")


            # ---- data base set ----
            # Auto remark dataset load
#            csv_url = "https://raw.githubusercontent.com/MSL-Official/inhouse_test/main/remarkbase.csv?token=GHSAT0AAAAAACDZUBCSDPJVR72QZW4BMHTAZEJIE2Q"
#            df_remarkbase = pd.read_csv(csv_url)
            csv_url_1 = [
                            ['IMPA', 'CONTENT', 'REMARK'],
                            ['', '200A', '200A'],
                            ['', '250A', '250A'],
                            ['', 'AIRLESS', 'AIRLESS'],
                            ['', 'ANODE', 'Anode'],
                            ['', 'MAINTENANCE FREE', 'Battery'],
                            ['', 'BREATHING AIR COMPRESSOR', 'BA Comp'],
                            ['', 'CARGO NET', 'CARGO NET w/cert - Vancouver need cert'],
                            ['', 'OVER LASHING', 'OVER LASHING CHAIN w/Canada cert - Need Certificate for Canadian Regulations'],
                            ['', 'FIRE HOSE', 'Cert check'],
                            ['', 'TEST LOAD WEIGHT', 'Cert check'],
                            ['', 'BARGE CHARGE', 'CHARGE'],
                            ['', 'BOAT CHARGE', 'CHARGE'],
                            ['', 'CUSTOMS CLEARANCE CHARGE', 'CHARGE'],
                            ['', 'DELIVERY CHARGE', 'CHARGE'],
                            ['', 'HATCH COVER CHAIN', 'Hatch Chain'],
                            ['', 'NO STOCK', 'Check'],
                            ['', 'SAME AS', 'Check'],
                            ['', 'FOG NOZZLE', 'Cert check'],
                            ['', 'FOG NOZZLES', 'Cert check'],
                            ['', 'AIRLESS SPRAY GUN IWATA', 'compatibility : CN TLS - JP IWATA - KR A-TECH HANDOK - HASCO - GRACO'],
                            ['', 'HIGH PRESSURE PAINT SPRAY', 'compatibility : CN TLS - JP IWATA - KR A-TECH HANDOK - HASCO - GRACO'],
                            ['', 'KYORITSU', 'KYORITSU logbook'],
                            ['', 'DIAMPHRAGM PUMP', 'Diaphragm pump'],
                            ['', 'WELDEN', 'Diaphragm pump'],
                            ['', 'WILDEN', 'Diaphragm pump'],
                            ['', 'EEBD', 'EEBD'],
                            ['', 'ENGINE STATIONERY', 'ENG-STA'],
                            ['', 'FIRE HOSE BOX', 'Fire Hose Box - Check w/Nozzle Bracket or not'],
                            ['', 'GANGWAY NET', 'GANGWAY NET - need cert in canada'],
                            ['', 'GANGWAY SAFETY NET', 'GANGWAY NET - need cert in canada'],
                            ['', 'GRAB CLOSING', 'GRAB CLOSING Wire - only one sheet(Long length) class cert with MTC for total length / only tag with ID number punching'],
                            ['', 'GRAB-CLOSING', 'GRAB CLOSING Wire - only one sheet(Long length) class cert with MTC for total length / only tag with ID number punching'],
                            ['', 'HAWSER', 'Hawser'],
                            ['', 'MOORING', 'Hawser'],
                            ['', 'TAIL ROPE', 'tail rope'],
                            ['', 'HAWSER PROTECTIVE', 'Protector'],
                            ['', 'LINE PROTECT', 'Protector'],
                            ['', 'LINE COVER', 'Protector'],
                            ['', 'PROTECTION COVER', 'Protector'],
                            ['', 'HOISTING WIRE', 'Hoisting Wire - Original cert(not long length / need one by one) / Cert number ID Punch / steel tag / Various port cert'],
                            ['', 'HOISTING WR', 'Hoisting Wire - Original cert(not long length / need one by one) / Cert number ID Punch / steel tag / Various port cert'],
                            ['', 'LUFFING WIRE', 'Luffing Wire - Original cert(not long length / need one by one) / Cert number ID Punch / steel tag / Various port cert'],
                            ['', 'LUFFING WR', 'Luffing Wire - Original cert(not long length / need one by one) / Cert number ID Punch / steel tag / Various port cert'],
                            ['', 'STEAM HOSE', 'SHOULD NO BE CUT'],
                            ['', 'INTRINSIC', 'INTRINSICALLY Safe'],
                            ['', 'INTRINSICALLY', 'INTRINSICALLY Safe'],
                            ['331097', 'CARTRIDGE TYPE CAPSULES FOR SMOKE & HEAT DETECT', 'Maker check'],
                            ['', 'BALNITURATE', 'Narcotic check'],
                            ['', 'CODEINE', 'Narcotic check'],
                            ['', 'CODINE', 'Narcotic check'],
                            ['', 'DIAZEPAM', 'Narcotic check'],
                            ['', 'MIDAZOLAM', 'Narcotic check'],
                            ['', 'MORPHINE', 'Narcotic check'],
                            ['', 'NITRAZEPAM', 'Narcotic check'],
                            ['', 'PENTAZOCINE', 'Narcotic check'],
                            ['', 'PERMAGANATE', 'Narcotic check'],
                            ['', 'PREDNISOLONE', 'Narcotic check'],
                            ['', 'PERMANGANATE', 'Narcotic check'],
                            ['', 'PHENOBARBITAL', 'Narcotic check'],
                            ['', 'POTASSIUM', 'Narcotic check'],
                            ['', 'TETANUS', 'Narcotic check'],
                            ['', 'TRAMADOL', 'Narcotic check'],
                            ['', 'LIGNOCAINE', 'Narcotic check'],
                            ['391716', 'ANTIMALARIAL PROPHYLAXIS', 'Need when vsl calls Southeast Asia'],
                            ['', 'SUEZ SEARCH', 'Need cert HK,CE,etc'],
                            ['', 'MARITEC', 'Owner arrange'],
                            ['', 'OIL SAMPLING', 'Owner arrange'],
                            ['', 'CUBITINNER', 'Owner arrange'],
                            ['', 'ACETYLENE (A-40)', 'Owner arrange'],
                            ['', 'ALKALINITY TABLET', 'Owner arrange'],
                            ['', 'AUTOTREAT', 'Owner arrange'],
                            ['', 'BOILERMATE', 'Owner arrange'],
                            ['', 'DISCLEAN', 'Owner arrange'],
                            ['', 'DIESELGUARD', 'Owner arrange'],
                            ['', 'ELECTROLYTE', 'Owner arrange'],
                            ['', 'ENVIROCLEAN', 'Owner arrange'],
                            ['', 'GAS DETECTOR', 'Gas detector - 4 kinds Gas Bulk Combustible gas is almost methane / O2, CO, H2S is confirmed , CH4-Methane (HC isobutane is rare)'],
                            ['', 'OWNER ARRANGE', 'Owner arrange'],
                            ['', 'OWNER ARRANGEMENT', 'Owner arrange'],
                            ['', 'OWNER ARRANGMENT', 'Owner arrange'],
                            ['', 'OXYGEN (O-40)', 'Owner arrange'],
                            ['', 'OXYGEN CONTROL', 'Owner arrange'],
                            ['', 'UNITOR', 'Owner arrange'],
                            ['', 'NOON REPORT', 'Owner arrange / Use vessel Excel File ?'],
                            ['', 'PILOT LADDER', 'PILOT LADDER - Need 2pc of spec-plate / Mumbai, india need product within 2 years and cert / Plate : fixed underside the steps. One on the upper part and the other one at the lower / ISO799-2019 each step with a mechanical clamping device type / Lower Manila rope only cut(not roop) / handrail manila rope need 28-32mm (both beside of pilot ladder)'],
                            ['', 'CARBON BRUSH', 'Please provide to us the Model/Type/Picture'],
                            ['', 'OXYGEN UNIT', 'Portable Oxygen - Need Bar check / CAN\'T DISPATCH BY SEA / 40Lx200BAR can be sucked by 2 people at the same time Fixed item + 2Lx200BAR portable type + 2Lx200BAR 1 piece portable total 8800L'],
                            ['', 'OXYGEN CYLINDER', 'Portable Oxygen - Need Bar check / CAN\'T DISPATCH BY SEA / 40Lx200BAR can be sucked by 2 people at the same time Fixed item + 2Lx200BAR portable type + 2Lx200BAR 1 piece portable total 8800L'],
                            ['', 'SPARE OXYGEN CYLINDER', 'Portable Oxygen - Need Bar check / CAN\'T DISPATCH BY SEA / 40Lx200BAR can be sucked by 2 people at the same time Fixed item + 2Lx200BAR portable type + 2Lx200BAR 1 piece portable total 8800L'],
                            ['', 'DIN.', 'Thread DIN'],
                            ['', 'RAT POISON', 'Restricted in SGP'],
                            ['', 'SCBA', 'SCBA'],
                            ['', 'GARBAGE RECORD', 'Both Part 1 and Part 2 of SGP flag & Marshall flag exist separately'],
                            ['', 'OFFICIAL LOG BOOK ', 'SGP flag - MPA OFFICIAL LOGBOOK'],
                            ['', 'SPRAY', 'check seafreight'],
                            ['', 'STORM', 'Storm valve - w/cert'],
                            ['', 'TRANSCIEVER', 'TRANSCEIVER'],
                            ['', 'TRANSCEIVER', 'TRANSCEIVER'],
                            ['', 'HANDHELD MARINE RADIO', 'TRANSCEIVER']
            ]
            csv_url_2 = [
                            ['IMPA', 'CONTENT', 'REMARK'],
                            ['', 'IC-', 'TRANSCEIVER'],
                            ['', 'ICF', 'TRANSCEIVER'],
                            ['', 'ICOM', 'TRANSCEIVER'],
                            ['', 'PORTABE UHF RADIO', 'TRANSCEIVER'],
                            ['', 'CREW ROLL BOOK', 'panam flag - need w/Panam stamp'],
                            ['370871', 'OFFICIAL LOG BOOK ', 'panam flag - need w/Panam stamp'],
                            ['', 'OIL RECORD BOOK', 'panam flag - need w/Panam stamp'],
                            ['', 'TREADMIL', 'Walfare'],
                            ['', 'HIGH PRESSURE WASH', 'H.P. Washer'],
                            ['', 'HIGH PRESSURE WATER', 'H.P. Washer'],
                            ['336001', 'FIRE CONTROL PLAN', 'Fire Plan Symbol - need new 336092,300X400MM'],
                            ['', 'TEST GAS FOR SMOKE DETECTOR', 'seafreight check'],
                            ['', 'TEST SPRAY', 'seafreight check'],
                            ['', 'AEROSOL', 'seafreight check'],
                            ['', 'AIR FRESHENER', 'seafreight check'],
                            ['', 'AIR REFRESHENER', 'seafreight check'],
                            ['', 'AIR SALONPASS', 'seafreight check'],
                            ['', 'BUTANE FUEL', 'seafreight check'],
                            ['', 'CAN\'T DESPATCH BY SEA', 'seafreight check'],
                            ['617017', 'COMPACT GAS TORCH', 'seafreight check'],
                            ['617017', 'COMPACT GAS TORCHES SPARE GAS CARTRIDGE', 'seafreight check'],
                            ['617017', 'SPARE GAS CARTRIDGE', 'seafreight check'],
                            ['795506', 'ELECTRIC CONTACT RESTORER', 'seafreight check'],
                            ['', 'ETHYLENE GLYCOL', 'seafreight check'],
                            ['', 'HYDRAULIC ACID', 'seafreight check'],
                            ['', 'HYDROCHLORIC ACID', 'seafreight check'],
                            ['', 'INSECT KILLER', 'seafreight check'],
                            ['', 'INSECTICIDE SPRAY', 'seafreight check'],
                            ['', 'GASOLINE', 'seafreight check'],
                            ['450107', 'KEROSENE', 'seafreight check'],
                            ['', 'LEAK DETECTOR', 'seafreight check'],
                            ['', 'MICRO CHECK', 'seafreight check'],
                            ['', 'METAL POLISH', 'seafreight check'],
                            ['', 'MACKINRY', 'seafreight check'],
                            ['', 'MERCURY', 'Delivery to Japan and Arrangements in Japan not possible'],
                            ['', 'PENETRAING OIL', 'seafreight check'],
                            ['', 'PENETRATING OIL', 'seafreight check'],
                            ['', 'SPHYGMOMANOMETER', 'Mercury check'],
                            ['', 'WD 40', 'seafreight check'],
                            ['', 'WD40', 'seafreight check'],
                            ['', 'WD-40', 'seafreight check'],
                            ['', 'ARC WELDER', 'ARC WELDER'],
                            ['', 'LED FLUORESCENT', 'Check double-side power supply or single-side power supply'],
                            ['', 'BRAZIL', 'Flag Double-sided printing'],
                            ['', 'BRASIL', 'Flag Double-sided printing'],
                            ['', 'BRAZILIAN', 'Flag Double-sided printing']
            ]
            df_remarkbase_1 = pd.DataFrame(csv_url_1[1:], columns=csv_url_1[0])
            df_remarkbase_2 = pd.DataFrame(csv_url_2[1:], columns=csv_url_2[0])

            
            # ---- start main perocess ----
            # upload process
            with st.form("upload-form", clear_on_submit=True):
                uploaded_files = st.file_uploader("", type=["xlsx", "xls","docx","csv"], accept_multiple_files=True)
                submitted = st.form_submit_button("Click to convert")
                if submitted:
                    for file in uploaded_files:
                        st.caption(file.name)
   
            # data processing
            if submitted and uploaded_files:
                prog_info.title("Proceed convert process")
                time_sta = tm.time()
                with st.spinner('Processing'):
                    # Content process
                    file_name = None
                    vessel_name = ""
                    combined_data = []
                    for file in uploaded_files:
                        # File name shape
                        file_name = filenameshape(file)

                        # Advance preparation / file control
                        sheet_data, file_type = data_sorting(uploaded_files, file, all_sht_flag)

                        # Converter main process / content control
                        combined_data, vessel_name, port_name = process_data(sheet_data, combined_data, file_type, file_name, order, page_col, vessel_name, mode)


                    # ---- process after data combined----
                    # Header insert
                    combined_data.insert(0,[None] * len(combined_data[0])) # add empty row

                    # all "" to be NULL
                    for row in combined_data:
                        for i in range(len(row)):
                            if row[i] == "":
                                row[i] = None
#                    st.write(combined_data)####
                    tim = tm.time() - time_sta
                    st.sidebar.caption("Data sorted & Combined : " + str(round(tim,3)) + 'sec')


                    # ---- Auro remark process ----
                    remark_content = []
                    if auto_remark_flag:
                        combined_data, remark_content = autoremark(combined_data, df_remarkbase_1, df_remarkbase_2, order)
#                    st.write(combined_data)####
                    tim = tm.time() - time_sta
                    st.sidebar.caption("Auto remark finished : " + str(round(tim,3)) + 'sec')


                    # ---- Excel file control & conbined data refflect ----
                    workbook = Workbook()
                    sheet = workbook.active
                    for row in combined_data:
                        sheet.append(row)
                    sheet.title = "Template"
                    sheet.sheet_view.zoomScale = zoom
                    sheet = excelcontrol(combined_data, sheet, auto_remark_flag, remark_content, mode)
                    
                    # vessen info process
                    vsl_info = ""
                    if vessel_name:
                        vsl_info = vessel_name.strip()
                    if port_name:
                        vsl_info = str(vsl_info) + "," + str(port_name)
                    if vsl_info:
                        sheet['N2'] = vsl_info
#                    st.write(vsl_info)
                    
                    # workbook close process
                    workbook.save("Converted_excel_file.xlsx")
                    workbook.close()

                    tim = tm.time() - time_sta
                    st.sidebar.caption("File form fixed : " + str(round(tim,3)) + 'sec')


                    # ---- Prepare for download ----
                    # excel control
                    workbook.save("Converted_excel_file.xlsx")
                    workbook.close()
                    with open("Converted_excel_file.xlsx", "rb") as file:
                        file_data = file.read()
                        prog_info.title("File available for download")
                    tim = tm.time() - time_sta
                    st.sidebar.caption("Total running time : " + str(round(tim,3)) + 'sec')
                    
                    # start download
                    st.title("")
                    file_name = vessel_name + str(datetime.now().strftime("%Y%m%d")) + " Converted.xlsx"                   
                    if combined_data is not None:
                        download_button_clicked = st.download_button(
                            label = "**Download here**",
                            data = file_data,
                            file_name = file_name,
                        )

                        # convert info
                        st.code(file_name,language="visualBasic (visual-basic)")

                        # end process - clear all
                        os.remove("Converted_excel_file.xlsx")
                        st.cache_data.clear()
                        combined_data = []
                        file_data = None
                        if download_button_clicked:
                            st.experimental_rerun()


        # ---- pdf convert
        with tab2:
            # ---- start main perocess ----
            # advance setting
            list1 = ['','','','','','','','']
            df_x = pd.DataFrame([list1])
            df_x.columns = ['page', 'word', 'x1','x2','y1','y2','width','hight']
            int_page = 0
            ii_index = 0
            df_s_x = []
            file_data = None

            # Upload process
            with st.form("pdf-form", clear_on_submit=True):
                uploaded_filespdf = st.file_uploader("", type=["pdf",'jpg'], accept_multiple_files=True)
                submittedpdf = st.form_submit_button("Click to convert")
                if submittedpdf:
                    for file in uploaded_filespdf:
                        st.caption(file.name)

            # ---- data processing ----
            if submittedpdf and uploaded_filespdf:
                time_stapdf = tm.time()
                with st.spinner('Processing'):
                    for file in uploaded_filespdf:
                        if str(file.name).lower().endswith("pdf"):
                            # ---- pdf open & coordinate x,y----
#                            work_file = 'work_file.xlsx'
                            # pdfminer3 setting
                            laparams = LAParams(line_overlap=0.1,
                                                word_margin=0.1,
                                                char_margin=0.1,
                                                line_margin=0.1,
                                                detect_vertical=True)
                            resource_manager = PDFResourceManager()
                            device = PDFPageAggregator(resource_manager, laparams=laparams)
                            interpreter = PDFPageInterpreter(resource_manager, device)

                            # read pdf with coordinate x,y
                            for page in PDFPage.get_pages(file):
                                int_page = int_page + 1
                                interpreter.process_page(page)
                                layout = device.get_result()
                                for lt in layout:
                                    # Standard output only for TTextContainer
                                    if isinstance(lt, LTTextContainer):
                                        df_x.loc[ii_index] = [int_page,'{}'.format(lt.get_text().strip()), lt.x0 , lt.x1 ,\
                                        841 - lt.y0 + (int_page - 1) * 841,841 - lt.y1  + (int_page - 1) * 841,lt.width ,lt.height ]
                                        ii_index = ii_index + 1
                            device.close()
#                            st.write(df_x)####

                            if df_x.empty:
                                # sort by x1
                                df_s_x = df_x.sort_values(['x1','y2'], ascending=[True,True])
                                # Calculate vertical pitch
                                h_min = 100
                                for i in range(len(df_s_x)):
                                    if i > 0:
                                        if df_s_x.iloc[i-1,2] == df_s_x.iloc[i,2]:
                                            h_sa = df_s_x.iloc[i,5] - df_s_x.iloc[i-1,5] 
                                            if h_sa > 1.0 and h_min > h_sa:
                                                h_min = h_sa

                                # ---- sort by excel with coordinated x,y ----
                                # set workbook
                                workbook = Workbook()
                                sheet = workbook.active
                                sheet.title = "Template"

                                # place data fm sorted df
                                j = 1
                                width_x = 0
                                for i in range(len(df_s_x)):
                                    y = df_s_x.iloc[i,5] // (math.ceil(h_min*10)/10) + 1
                                    c1 = sheet.cell(row=int(y), column=j)
                                    if c1.value == None:
                                        c1.value = df_s_x.iloc[i,1]
                                    else:
                                        # adjust the column width
                                        sheet.column_dimensions[sheet.cell(row=1, column=j).column_letter].width = (df_s_x.iloc[i,2]/5.98- width_x )
                                        width_x = df_s_x.iloc[i,2]/5.98
                                        j = j + 1
                                        c1 = sheet.cell(row=int(y), column=j)
                                        c1.value = df_s_x.iloc[i,1]

                                # delete empty row with 5rows
                                empty_row_count = 0
                                max_empty_rows = 5
                                rows_to_delete = []
                                for row in sheet.iter_rows():
                                    is_empty_row = all(cell.value is None for cell in row)
                                    if is_empty_row:
                                        empty_row_count += 1
                                        if empty_row_count >= max_empty_rows:
                                            rows_to_delete.append(row[0].row)
                                    else:
                                        empty_row_count = 0
                                        
                                for row_index in reversed(rows_to_delete):
                                     sheet.delete_rows(row_index, 1)
                    

                                # ---- Prepare for download ----
                                # excel control
                                workbook.save("Converted_excel_file.xlsx")
                                workbook.close()
                                with open("Converted_excel_file.xlsx", "rb") as file:
                                    file_data = file.read()
                                    prog_info.title("File available for download")
                                timpdf = tm.time() - time_stapdf
                                st.sidebar.caption("Total running time : " + str(round(timpdf,3)) + 'sec')
                                
                                # start download    
                                st.title("")
                                file_name = str(datetime.now().strftime("%Y%m%d")) + " Converted fm PDF.xlsx"                   
                                if file_data is not None:
                                    download_button_clicked_pdf = st.download_button(
                                        label = "**Download here**",
                                        data = file_data,
                                        file_name = file_name,
                                    )
                                    # convert info
                                    st.code(file_name,language="visualBasic (visual-basic)")

                                    # end process - clear all
                                    os.remove("Converted_excel_file.xlsx")
                                    st.cache_data.clear()
                                    df_s_x = None
                                    df_x = None
                                    file_data = None
                                    if download_button_clicked_pdf:
                                        st.experimental_rerun()


                            # if not str embedding pdf                    
                            else:
                                st.write("not str embedding pdf")
#                                st.stop()####
                                images = pdf2image.convert_from_bytes(file.read())
                                for page in images:
                                    st.image(page, use_column_width=True)




                                reader = easyocr.Reader(['en'],gpu = False)
                                results = reader.readtext(file.getvalue(), detail = 0)
                                for result in results:
                                    st.write(result[1], " ",result[2])



        # ---- online manual ----
        with tab3:
            # ---- manual ----
            st.subheader("Converter Manual - How to use")
            st.write("1. file select")
            st.write("ファイルをドラッグアンドドロップすると処理が始まります-複数ファイルも可能です / Browse fileでの選択も可能です")
            st.write("Drag and drop files to start processing- Multiple files are possible / We can also select by Browse file button")
            st.write(" ")
            st.write("2. Mode/Country select")
            st.write("出力時のフォーマットを選択します / mashin system store mode等")
            st.write("Select output format / mashin system store mode etc.")
            st.write(" ")
            st.write("3. Sheet process select")
            st.write("変換対象 -開いた時のsheet- or -全シート- を選択")
            st.write("Select conversion target -sheet when opened- or -all sheets-")
            st.write(" ")
            st.write("4. Auto remark process select")
            st.write("Auto remarkを適用可否選択 -No use- or -Apply-")
            st.write("Select whether to apply Auto remark -No use- or -Apply-")
            st.write(" ")
            st.write("5. file download")
            st.write("処理か完了すると -Download Excel-ボタンが出現します / 押すとファイルのファウンロードを行います")
            st.write("-Download Excel- button will appear when processing is complete / Press to download the file")


        # ---- tips ----
        with tab4:
            # ---- tips ----
            st.subheader("Tips - Augmented Recognition")
            st.write("・変換されないファイルがある場合は福岡 海老原宛てに御連絡下さい / If there are files that cannot be converted, please contact Fukuoka-office Y.Ebihara(Mr.).")
            st.write("・ConverterはArticleやQTY等 各項目の位置や座標特定を試みます / Converter tries to identify the position and coordinates of each item such as Article and QTY.")
            st.write("・Article特定不可時の救済策として、Article行を手入力する機構を表示します / Display the mechanism to manually enter the Article row as a remedy when the Article cannot be specified.")
            st.write("・Excelファイルの場合、シートのロックやセル結合等は強制的に解除しています / For Excel files, sheet locks and cell merging are forcibly released.")
            st.write("・QTYとUnitが結合している場合、分離後します。元からUnitが抜けているケースにご注意下さい / If QTY and Unit are combined, after separation. Please pay attention to the case where Unit is missing from the original.")
            st.write("・Auto-remarkはArticleから要確認事項を教えてくれたり、dept内重複を検出してくれます / Auto-remark will tell you what to check from Article and detect duplication in dept.")
            st.write("・もし、dept-columnが無い場合、Articleからdept内容特定を試み、それでも不明な場合、sheet_nameやfile_nameからも特定を試みます / If there is no dept-column, try to identify the dept content from Article, and if still unknown, try to identify from sheet_name and file_name.")
            st.write("・Excelのxls fileは変換可能ですが、長い時間を要します / Excel xls file can be converted but it takes a long time")
            st.write("・Word Docxは変換可能ですが、Doc fileは未対応です-考え中です / Word Docx can be converted, but doc file is not - I'm thinking about it...")
            st.write("・Converterと別tabで、文字埋込PDF -> Excel file変換を行う事が出来ます。精度にばらつきあり / Character embedding PDF -> Excel file conversion can be performed in a separate tab from Converter. Accuracy varies")
            st.write("・もし、隠しシートも強制的に読み取ってしまったら、それは興味深い出来事ですね / If you force the hidden sheet to be read, that's an interesting event.")
            st.write("・Web pageの動きが重い？そんな時はcache clearだ! / Is the web page slow to move? In such a case, Go cache clear!")


    # ---- Pass-code error process ----
    elif code:
        st.error("Invalid code.")
        code_placeholder.empty() 
        user_placeholder.empty()
        st.stop()


if __name__ == "__main__":
    main()